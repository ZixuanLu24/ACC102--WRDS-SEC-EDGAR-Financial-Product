import streamlit as st
import pandas as pd
import numpy as np
import datetime
import io
import os
import re
import uuid
import requests
import json
from sqlalchemy import create_engine, text
import plotly.graph_objects as go
import plotly.express as px
import openpyxl
from openpyxl.drawing.image import Image as OpenpyxlImageExcel

# docx imports, adding OxmlElement and qn for underlying XML tables and layout
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

# =========================================================
# Export Safety Utilities (Never Crash Export)
# =========================================================
def safe_sheet_name(name: str) -> str:
    """Excel sheet name must be <= 31 chars and cannot contain : \ / ? * [ ]"""
    if name is None:
        return "Sheet"
    name = str(name)
    name = re.sub(r'[:\\/?*\[\]]', "_", name)
    name = name.strip()
    if len(name) == 0:
        return "Sheet"
    return name[:31]

def safe_to_excel(writer, df: pd.DataFrame, sheet_name: str, index: bool = False):
    """Safely write DataFrame to Excel. Skips empty/None df."""
    if df is None:
        return False
    if isinstance(df, pd.DataFrame) and df.empty:
        return False
    try:
        df.to_excel(writer, sheet_name=safe_sheet_name(sheet_name), index=index)
        return True
    except Exception:
        return False

def safe_plotly_to_png_bytes(fig):
    """Convert Plotly figure to PNG bytes safely. Returns None if conversion fails."""
    if fig is None:
        return None
    try:
        return fig.to_image(format="png", engine="kaleido")
    except Exception:
        try:
            return fig.to_image(format="png")
        except Exception:
            return None

def safe_add_image_to_sheet(ws, fig, cell="A1", temp_files=None, width=None, height=None):
    """
    Safely add Plotly figure as image into an openpyxl worksheet.
    Defers deletion of the temporary file until after the workbook is fully saved.
    """
    img_bytes = safe_plotly_to_png_bytes(fig)
    if img_bytes is None:
        return False

    try:
        tmp_name = f"_tmp_chart_{uuid.uuid4().hex}.png"
        with open(tmp_name, "wb") as f:
            f.write(img_bytes)

        img = OpenpyxlImageExcel(tmp_name)
        if width: img.width = width
        if height: img.height = height
        ws.add_image(img, cell)

        # Do not delete immediately; openpyxl needs the file during workbook save.
        if temp_files is not None:
            temp_files.append(tmp_name)

        return True
    except Exception:
        return False


# Ensure Plotly line colors are visible on dark Excel theme (avoid black lines)
px.defaults.template = "plotly_dark"
px.defaults.color_discrete_sequence = px.colors.qualitative.Plotly

# =========================================================
# Page Config
# =========================================================
st.set_page_config(
    page_title="WRDS Multi-Dimensional Financial Analysis",
    layout="wide",
    page_icon="🏦"
)

# =========================================================
# Global Style
# =========================================================
st.markdown("""
<style>
html, body, [class*="css"]  {
    background-color: #0b0f14 !important;
    color: #e0e0e0 !important;
    font-family: "Segoe UI", Arial, sans-serif;
}

section[data-testid="stSidebar"] {
    background-color: #0e131a !important;
    border-right: 1px solid #1f2a38;
}

.block-container {
    padding-top: 1.0rem;
    padding-bottom: 2rem;
}

div[data-testid="stMetric"] {
    background-color: #121923;
    border: 1px solid #1f2a38;
    border-radius: 12px;
    padding: 14px;
}

div[data-testid="stMetricLabel"] {
    color: #b7c3d0 !important;
    font-size: 13px !important;
}

div[data-testid="stMetricValue"] {
    color: #00ff99 !important;
    font-size: 24px !important;
    font-weight: 600 !important;
}

div[data-testid="stMetricDelta"] {
    color: #f0b90b !important;
}

.stButton > button {
    background-color: #0b5ed7;
    color: white;
    border-radius: 10px;
    border: 1px solid #1f2a38;
    font-weight: 600;
    padding: 0.5rem 1rem;
}

.stButton > button:hover {
    background-color: #084298;
    border-color: #2b3a50;
}

div[data-testid="stDataFrame"] {
    background-color: #0e131a !important;
    border: 1px solid #1f2a38;
    border-radius: 12px;
}

div[data-testid="stExpander"] {
    background-color: #0e131a !important;
    border: 1px solid #1f2a38;
    border-radius: 12px;
}

hr {
    border: none;
    border-top: 1px solid #1f2a38;
}

h1, h2, h3, h4 {
    color: #ffffff !important;
}

small, p, span, label {
    color: #c9d1d9 !important;
}

code {
    background-color: #121923 !important;
    color: #00ff99 !important;
    border-radius: 6px;
    padding: 2px 6px;
}

.stAlert > div {
    border-radius: 12px;
    border: 1px solid #1f2a38;
}
</style>
""", unsafe_allow_html=True)

# =========================================================
# Metric Reference
# =========================================================
METRIC_REFERENCE_DATA = [
    {
        "Metric": "daily_return_py",
        "English Name": "Daily Return",
        "Formula": "(P_t / P_{t-1}) - 1",
        "Interpretation": "Positive means gain, negative means loss over one day.",
        "Common Reference Range": "For many large-cap stocks, daily moves often fall within about -2% to +2%; volatile stocks may exceed ±5%.",
        "Open-source / Historical Reference": "Commonly used in financial textbooks and market analytics practice.",
        "How to Judge": "Higher is not always better on a single day; interpret over a period together with volatility."
    },
    {
        "Metric": "log_return_py",
        "English Name": "Log Return",
        "Formula": "ln(P_t / P_{t-1})",
        "Interpretation": "Approximately equal to simple return when price changes are small.",
        "Common Reference Range": "Usually close to daily return in magnitude for normal market moves.",
        "Open-source / Historical Reference": "Widely used in econometrics and quantitative finance.",
        "How to Judge": "Better suited for modeling and statistical analysis."
    },
    {
        "Metric": "annualized_return_py",
        "English Name": "Annualized Return",
        "Formula": "(1 + Total Return)^(TradingDays / N) - 1",
        "Interpretation": "Represents annual-equivalent compounded return.",
        "Common Reference Range": "Broad equity markets are often cited around 6%–12% in long-run history, depending on market and sample period.",
        "Open-source / Historical Reference": "Standard measure in investment analysis and portfolio reporting.",
        "How to Judge": "Higher is generally better, but must be evaluated with risk metrics."
    },
    {
        "Metric": "cumulative_return_py",
        "English Name": "Cumulative Return",
        "Formula": "∏(1+r_t) - 1",
        "Interpretation": "Shows total performance over the selected horizon.",
        "Common Reference Range": "No fixed universal standard; depends on date window.",
        "Open-source / Historical Reference": "Standard performance metric in portfolio and benchmark comparison.",
        "How to Judge": "Useful for comparing assets over the same date range."
    },
    {
        "Metric": "excess_return_py",
        "English Name": "Excess Return",
        "Formula": "R_excess = R_asset - R_benchmark",
        "Interpretation": "Positive means outperformance versus benchmark.",
        "Common Reference Range": "Positive annualized excess return is generally desirable.",
        "Open-source / Historical Reference": "Widely used in active management and CAPM-style analysis.",
        "How to Judge": "Positive indicates relative outperformance; benchmark selection matters."
    },
    {
        "Metric": "volatility_daily_py",
        "English Name": "Daily Volatility",
        "Formula": "σ_d = StdDev(r_t)",
        "Interpretation": "Higher values indicate greater day-to-day fluctuation.",
        "Common Reference Range": "For many large-cap stocks, around 0.5%–2% daily volatility may be common.",
        "Open-source / Historical Reference": "Core metric in risk analysis and portfolio theory.",
        "How to Judge": "Lower means more stable, but not necessarily higher return."
    },
    {
        "Metric": "volatility_annualized_py",
        "English Name": "Annualized Volatility",
        "Formula": "σ_a = σ_d · √(TradingDays)",
        "Interpretation": "Measures annualized risk level.",
        "Common Reference Range": "~10%–15% low/moderate, ~15%–25% medium, ~25%+ high for equities.",
        "Open-source / Historical Reference": "A core risk metric used in portfolio reports and public fact sheets.",
        "How to Judge": "Lower is generally preferred for the same level of return."
    },
    {
        "Metric": "var_hist_95_py",
        "English Name": "VaR 95% (Historical)",
        "Formula": "-Percentile(Returns, 5%)",
        "Interpretation": "At 95% confidence, daily loss is expected not to exceed this threshold on most days.",
        "Common Reference Range": "For many equities, daily 95% VaR around 1%–3% may be common.",
        "Open-source / Historical Reference": "Widely used in academic and practical risk management.",
        "How to Judge": "Smaller magnitude is safer."
    },
    {
        "Metric": "var_hist_99_py",
        "English Name": "VaR 99% (Historical)",
        "Formula": "-Percentile(Returns, 1%)",
        "Interpretation": "Captures more extreme downside than 95% VaR.",
        "Common Reference Range": "Typically larger than 95% VaR in magnitude.",
        "Open-source / Historical Reference": "Standard extreme-risk measure in risk analytics.",
        "How to Judge": "Lower is safer."
    },
    {
        "Metric": "var_param_95_py",
        "English Name": "VaR 95% (Parametric)",
        "Formula": "-(μ + z_{0.95}·σ)",
        "Interpretation": "Model-based risk estimate under distribution assumptions.",
        "Common Reference Range": "Often close to historical VaR but may understate fat-tail risk.",
        "Open-source / Historical Reference": "Classical risk management approach in textbooks.",
        "How to Judge": "Useful for model-based estimation; compare with historical VaR."
    },
    {
        "Metric": "var_param_99_py",
        "English Name": "VaR 99% (Parametric)",
        "Formula": "-(μ + z_{0.99}·σ)",
        "Interpretation": "More conservative than 95% parametric VaR.",
        "Common Reference Range": "Should be larger than 95% parametric VaR in magnitude.",
        "Open-source / Historical Reference": "Common in quantitative finance and risk reports.",
        "How to Judge": "Higher means larger estimated tail risk."
    },
    {
        "Metric": "cvar_hist_95_py",
        "English Name": "CVaR 95% (Historical)",
        "Formula": "-Mean(Returns | Returns <= 5th percentile)",
        "Interpretation": "Average loss in the worst 5% of cases.",
        "Common Reference Range": "Must be equal to or larger than 95% VaR in magnitude.",
        "Open-source / Historical Reference": "Preferred tail-risk metric in advanced risk frameworks.",
        "How to Judge": "Lower is safer; a much larger CVaR than VaR indicates severe tail losses."
    },
    {
        "Metric": "cvar_hist_99_py",
        "English Name": "CVaR 99% (Historical)",
        "Formula": "-Mean(Returns | Returns <= 1st percentile)",
        "Interpretation": "Average loss in the worst 1% of cases.",
        "Common Reference Range": "Typically more conservative than CVaR 95%.",
        "Open-source / Historical Reference": "Standard tail-risk measure in advanced portfolio risk analysis.",
        "How to Judge": "Lower is better."
    },
    {
        "Metric": "max_drawdown_py",
        "English Name": "Maximum Drawdown",
        "Formula": "min((Wealth_t / RunningMax_t) - 1)",
        "Interpretation": "Measures the worst peak-to-trough loss.",
        "Common Reference Range": "For equities, drawdowns beyond -20% are often considered substantial; crisis periods may be much worse.",
        "Open-source / Historical Reference": "Widely used in fund reports and portfolio risk analysis.",
        "How to Judge": "Less negative is better."
    },
    {
        "Metric": "sharpe_ratio_py",
        "English Name": "Sharpe Ratio",
        "Formula": "Sharpe = (μ_a - r_f) / σ_a",
        "Interpretation": "Measures risk-adjusted return using total volatility.",
        "Common Reference Range": "<1 weak, 1–2 acceptable/good, 2–3 very good, >3 very strong (rule of thumb).",
        "Open-source / Historical Reference": "One of the most widely used performance ratios in portfolio management.",
        "How to Judge": "Higher is generally better."
    },
    {
        "Metric": "calmar_ratio_py",
        "English Name": "Calmar Ratio",
        "Formula": "Calmar = μ_a / |MDD|",
        "Interpretation": "Measures return relative to worst historical drawdown.",
        "Common Reference Range": ">1 often considered good; depends on asset class and period.",
        "Open-source / Historical Reference": "Common in hedge fund and alternative investment analysis.",
        "How to Judge": "Higher is better."
    },
    {
        "Metric": "beta_py",
        "English Name": "Beta",
        "Formula": "β = Cov(R_i, R_m) / Var(R_m)",
        "Interpretation": "Beta > 1 implies more sensitive than benchmark; Beta < 1 implies less sensitive.",
        "Common Reference Range": "~1 market-like, >1 high sensitivity, <1 defensive.",
        "Open-source / Historical Reference": "Core concept in CAPM and equity analysis.",
        "How to Judge": "Interpret relative to benchmark."
    },
    {
        "Metric": "information_ratio_py",
        "English Name": "Information Ratio",
        "Formula": "IR = (μ_p - μ_b) / TE",
        "Interpretation": "Measures consistency of excess return versus benchmark.",
        "Common Reference Range": ">0.5 often viewed as meaningful; >1 strong in many practical contexts.",
        "Open-source / Historical Reference": "Common in active management evaluation.",
        "How to Judge": "Higher is better."
    },
    {
        "Metric": "correlation_matrix",
        "English Name": "Correlation Matrix",
        "Formula": "ρ_{ij} = Cov(R_i, R_j) / (σ_i·σ_j)",
        "Interpretation": "Values range from -1 to 1; lower correlation may improve diversification.",
        "Common Reference Range": "0.8–1.0 high positive, 0.5–0.8 moderate-high positive, -0.2–0.2 weak relation, below -0.5 strong negative relation.",
        "Open-source / Historical Reference": "Standard in diversification and asset allocation analysis.",
        "How to Judge": "Lower or negative correlation is often better for diversification."
    },
    {
        "Metric": "rolling_mean_return_py",
        "English Name": "Rolling Mean Return",
        "Formula": "Mean(r_{t-n+1}, …, r_t)",
        "Interpretation": "Shows how short-to-medium term average return evolves over time.",
        "Common Reference Range": "No fixed standard; depends on window length and market state.",
        "Open-source / Historical Reference": "Common in time-varying performance analysis.",
        "How to Judge": "Higher may indicate improving return trend, but should be read with rolling volatility."
    },
    {
        "Metric": "rolling_volatility_py",
        "English Name": "Rolling Volatility",
        "Formula": "StdDev(r_{t-n+1}, …, r_t) · √(TradingDays)",
        "Interpretation": "Shows how risk changes over time.",
        "Common Reference Range": "No universal threshold; compare across firms or periods.",
        "Open-source / Historical Reference": "Standard in dynamic risk monitoring.",
        "How to Judge": "Higher means risk is rising over the selected rolling window."
    }
]

REFERENCE_DF = pd.DataFrame(METRIC_REFERENCE_DATA)

# =========================================================
# Mappings
# =========================================================
COLUMN_MAPPING = {
    'date': 'Trading Date',
    'permno': 'Permanent Security Identifier (PERMNO)',
    'permco': 'Permanent Company Identifier (PERMCO)',
    'hexcd': 'Exchange Code',
    'cusip': 'CUSIP Identifier',
    'comnam': 'Company Name',
    'ticker': 'Ticker Symbol',
    'prc': 'Closing Price / Bid-Ask Average',
    'vol': 'Trading Volume',
    'ret': 'Stock Return',
    'retx': 'Return Excluding Dividends',
    'bid': 'Bid Price',
    'ask': 'Ask Price',
    'bidlo': 'Lowest Bid Price',
    'askhi': 'Highest Ask Price',
    'spread': 'Bid-Ask Spread',
    'openprc': 'Opening Price',
    'numtrd': 'Number of Trades',
    'shrout': 'Shares Outstanding',
    'cfacpr': 'Cumulative Factor to Adjust Price',
    'cfacshr': 'Cumulative Factor to Adjust Shares',
    'dlret': 'Delisting Return',
    'dlstcd': 'Delisting Code',
    'altprc': 'Alternative Price',
    'divamt': 'Dividend Amount',
    'facpr': 'Price Adjustment Factor',
    'facshr': 'Share Adjustment Factor',
    'trdstat': 'Trading Status',
    'sprtrn': 'S&P Return',
    'nwperm': 'New PERMNO Indicator'
}

FIN_ANALYSIS_MAPPING = {
    'daily_return_py': 'Daily Return (Python-calculated from Price or WRDS ret)',
    'log_return_py': 'Log Return (Python-calculated)',
    'annualized_return_py': 'Annualized Return',
    'cumulative_return_py': 'Cumulative Return',
    'excess_return_py': 'Excess Return vs Benchmark/Mean',
    'volatility_daily_py': 'Daily Volatility',
    'volatility_annualized_py': 'Annualized Volatility',
    'var_hist_95_py': 'VaR 95% (Historical Simulation)',
    'var_hist_99_py': 'VaR 99% (Historical Simulation)',
    'var_param_95_py': 'VaR 95% (Parametric)',
    'var_param_99_py': 'VaR 99% (Parametric)',
    'cvar_hist_95_py': 'CVaR 95% (Historical)',
    'cvar_hist_99_py': 'CVaR 99% (Historical)',
    'max_drawdown_py': 'Maximum Drawdown',
    'sharpe_ratio_py': 'Sharpe Ratio',
    'calmar_ratio_py': 'Calmar Ratio',
    'beta_py': 'Beta',
    'treynor_ratio_py': 'Treynor Ratio',
    'information_ratio_py': 'Information Ratio',
    'rolling_mean_return_py': 'Rolling Mean Return',
    'rolling_volatility_py': 'Rolling Volatility'
}

# =========================================================
# Helpers
# =========================================================
def prettify_column_name(col_name):
    return COLUMN_MAPPING.get(col_name, col_name.replace('_', ' ').title())


def safe_numeric(series):
    return pd.to_numeric(series, errors='coerce')


def z_score_from_conf(conf_level):
    z_map = {
        0.90: 1.2815515655446004,
        0.95: 1.6448536269514722,
        0.975: 1.959963984540054,
        0.99: 2.3263478740408408
    }
    return z_map.get(conf_level, 1.6448536269514722)


def compute_max_drawdown(cumulative_return_series):
    if cumulative_return_series is None or cumulative_return_series.empty:
        return np.nan
    wealth = 1 + cumulative_return_series
    running_max = wealth.cummax()
    drawdown = wealth / running_max - 1
    return drawdown.min()


def compute_beta(asset_returns, benchmark_returns):
    merged = pd.concat([asset_returns, benchmark_returns], axis=1).dropna()
    if merged.shape[0] < 2:
        return np.nan

    x = pd.to_numeric(merged.iloc[:, 0], errors='coerce')
    y = pd.to_numeric(merged.iloc[:, 1], errors='coerce')
    merged2 = pd.concat([x, y], axis=1).dropna()

    if merged2.shape[0] < 2:
        return np.nan

    bench_var = merged2.iloc[:, 1].var(ddof=1)
    if pd.isna(bench_var) or bench_var == 0:
        return np.nan

    cov = np.cov(merged2.iloc[:, 0], merged2.iloc[:, 1], ddof=1)[0, 1]
    return cov / bench_var


def compute_tracking_error(active_returns, trading_days=252):
    s = pd.to_numeric(active_returns, errors='coerce').dropna()
    if len(s) < 2:
        return np.nan
    return s.std(ddof=1) * np.sqrt(trading_days)


def apply_price_adjustment(df, price_col='prc', factor_col='cfacpr', mode='multiply'):
    if df is None or df.empty:
        return pd.DataFrame()

    data = df.copy()

    if price_col not in data.columns:
        data['adj_prc'] = np.nan
        return data

    data[price_col] = pd.to_numeric(data[price_col], errors='coerce').abs()

    if factor_col not in data.columns:
        data['adj_prc'] = data[price_col]
        return data

    data[factor_col] = pd.to_numeric(data[factor_col], errors='coerce')

    if mode == 'divide':
        data['adj_prc'] = np.where(
            data[factor_col].notna() & (data[factor_col] != 0),
            data[price_col] / data[factor_col],
            np.nan
        )
    else:
        data['adj_prc'] = np.where(
            data[factor_col].notna(),
            data[price_col] * data[factor_col],
            np.nan
        )

    data['adj_prc'] = pd.to_numeric(data['adj_prc'], errors='coerce')
    return data


def detect_missing_dates(df, date_col='date', ticker_col='TICKER'):
    if df is None or df.empty or date_col not in df.columns or ticker_col not in df.columns:
        return pd.DataFrame(columns=[ticker_col, date_col, 'Issue'])

    tmp = df.copy()
    tmp[date_col] = pd.to_datetime(tmp[date_col], errors='coerce')
    tmp = tmp.dropna(subset=[ticker_col, date_col])

    issue_rows = []

    for ticker, group in tmp.groupby(ticker_col):
        group = group.sort_values(date_col)
        if group.empty:
            continue

        observed_dates = pd.to_datetime(group[date_col].unique())
        if len(observed_dates) == 0:
            continue

        full_range = pd.date_range(
            start=observed_dates.min(),
            end=observed_dates.max(),
            freq='B'
        )

        missing_dates = sorted(list(set(full_range) - set(observed_dates)))
        if len(missing_dates) > 0:
            issue_rows.append(pd.DataFrame({
                ticker_col: [ticker] * len(missing_dates),
                date_col: missing_dates,
                'Issue': ['Missing Trading Date'] * len(missing_dates)
            }))

    if issue_rows:
        return pd.concat(issue_rows, ignore_index=True)

    return pd.DataFrame(columns=[ticker_col, date_col, 'Issue'])


def clean_missing_and_anomalies(df, price_col='prc', volume_col='vol', return_col='ret', method='Forward Fill'):
    if df is None or df.empty:
        return pd.DataFrame(), pd.DataFrame(columns=['date', 'TICKER', 'Issue'])

    data = df.copy()
    issue_rows = []
    cleaned_parts = []

    for ticker, group in data.groupby('TICKER'):
        group = group.sort_values('date').copy()

        if price_col in group.columns:
            group[price_col] = pd.to_numeric(group[price_col], errors='coerce').abs()
            invalid_mask = group[price_col].isna() | (group[price_col] <= 0)
            if invalid_mask.any():
                issue_rows.append(pd.DataFrame({
                    'date': group.loc[invalid_mask, 'date'],
                    'TICKER': [ticker] * invalid_mask.sum(),
                    'Issue': 'Invalid or Missing Price'
                }))

        if volume_col and volume_col in group.columns:
            group[volume_col] = pd.to_numeric(group[volume_col], errors='coerce')
            invalid_mask = group[volume_col].isna() | (group[volume_col] < 0)
            if invalid_mask.any():
                issue_rows.append(pd.DataFrame({
                    'date': group.loc[invalid_mask, 'date'],
                    'TICKER': [ticker] * invalid_mask.sum(),
                    'Issue': 'Invalid or Missing Volume'
                }))

        if return_col and return_col in group.columns:
            group[return_col] = pd.to_numeric(group[return_col], errors='coerce')
            invalid_mask = group[return_col].isna()
            if invalid_mask.any():
                issue_rows.append(pd.DataFrame({
                    'date': group.loc[invalid_mask, 'date'],
                    'TICKER': [ticker] * invalid_mask.sum(),
                    'Issue': 'Missing Return'
                }))

        cols_candidate = [c for c in [price_col, volume_col, return_col, 'cfacpr'] if c and c in group.columns]

        if method == 'Forward Fill':
            if cols_candidate:
                group[cols_candidate] = group[cols_candidate].ffill()

        elif method == 'Mean Fill':
            for c in cols_candidate:
                mean_val = pd.to_numeric(group[c], errors='coerce').mean()
                group[c] = pd.to_numeric(group[c], errors='coerce').fillna(mean_val)

        elif method == 'Drop Missing':
            drop_cols = [c for c in [price_col, volume_col, return_col] if c and c in group.columns]
            if drop_cols:
                group = group.dropna(subset=drop_cols)

        if price_col in group.columns:
            group = group[group[price_col].notna() & (group[price_col] > 0)]

        if volume_col and volume_col in group.columns:
            group = group[group[volume_col].isna() | (group[volume_col] >= 0)]

        cleaned_parts.append(group)

    cleaned_df = (
        pd.concat(cleaned_parts, axis=0).sort_values(['date', 'TICKER'])
        if cleaned_parts else pd.DataFrame()
    )

    issues_df = (
        pd.concat(issue_rows, axis=0).sort_values(['date', 'TICKER'])
        if issue_rows else pd.DataFrame(columns=['date', 'TICKER', 'Issue'])
    )

    return cleaned_df, issues_df


def format_percentage(x, decimals=2):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"{x * 100:.{decimals}f}%"
    except Exception:
        return "N/A"


def format_float(x, decimals=4):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"{x:.{decimals}f}"
    except Exception:
        return "N/A"

def format_currency(val):
    if val is None or pd.isna(val): return "N/A"
    try:
        val = float(val)
        if abs(val) >= 1e9: return f"${val/1e9:,.2f} B"
        elif abs(val) >= 1e6: return f"${val/1e6:,.2f} M"
        else: return f"${val:,.2f}"
    except:
        return "N/A"

def sanitize_excel_filename(name):
    if not name:
        return "WRDS_Report"
    name = str(name).strip()
    name = re.sub(r'[\\/*?:\[\]]+', "_", name)
    name = re.sub(r'\s+', "_", name)
    return name[:120] if len(name) > 120 else name


@st.cache_data(ttl=3600)
def get_table_columns(_engine):
    try:
        query = """
        SELECT column_name
        FROM information_schema.columns
        WHERE table_schema = 'crsp'
          AND table_name = 'dsf'
        ORDER BY ordinal_position;
        """
        with _engine.connect() as conn:
            df = pd.read_sql(query, conn)

        raw_columns = df['column_name'].dropna().astype(str).str.strip().tolist()

        seen = set()
        db_columns = []
        for col in raw_columns:
            col_lower = col.lower()
            if col_lower not in seen:
                seen.add(col_lower)
                db_columns.append(col)

        display_options = []
        display_to_column = {}

        for col in db_columns:
            label = prettify_column_name(col)
            display_name = f"{col} ({label})"
            display_options.append(display_name)
            display_to_column[display_name] = col

        return db_columns, display_options, display_to_column

    except Exception as e:
        st.error(f"Failed to fetch columns: {e}")
        return [], [], {}


def build_long_format(metric_name, price_wide, returns_df, cumulative_return_df, rolling_vol_20_df):
    if metric_name == "Price":
        temp = price_wide.copy() if price_wide is not None else pd.DataFrame()
    elif metric_name == "Daily Return":
        temp = returns_df.copy() if returns_df is not None else pd.DataFrame()
    elif metric_name == "Cumulative Return":
        temp = cumulative_return_df.copy() if cumulative_return_df is not None else pd.DataFrame()
    elif metric_name == "Rolling Volatility (20D, Annualized)":
        temp = rolling_vol_20_df.copy() if rolling_vol_20_df is not None else pd.DataFrame()
    else:
        return pd.DataFrame(columns=["Date", "Ticker", "Value"])

    if temp.empty:
        return pd.DataFrame(columns=["Date", "Ticker", "Value"])

    base = temp.reset_index()
    long_df = base.melt(id_vars=base.columns[0], var_name="Ticker", value_name="Value")
    long_df.columns = ["Date", "Ticker", "Value"]
    long_df = long_df.dropna().sort_values(["Ticker", "Date"])
    return long_df


def build_long_format_from_df(df_wide):
    if df_wide is None or df_wide.empty:
        return pd.DataFrame(columns=["Date", "Ticker", "Value"])

    base = df_wide.reset_index()
    long_df = base.melt(id_vars=base.columns[0], var_name="Ticker", value_name="Value")
    long_df.columns = ["Date", "Ticker", "Value"]
    long_df = long_df.dropna().sort_values(["Ticker", "Date"])
    return long_df


def add_max_min_annotations(fig_obj, long_df):
    if fig_obj is None or long_df is None or long_df.empty:
        return fig_obj

    for ticker in long_df['Ticker'].dropna().unique():
        sub = long_df[long_df['Ticker'] == ticker].dropna(subset=['Value'])
        if sub.empty:
            continue

        max_idx = sub['Value'].idxmax()
        min_idx = sub['Value'].idxmin()

        max_row = sub.loc[max_idx]
        min_row = sub.loc[min_idx]

        fig_obj.add_trace(go.Scatter(
            x=[max_row['Date']],
            y=[max_row['Value']],
            mode='markers+text',
            text=['Max'],
            textposition='top center',
            marker=dict(size=9, symbol='diamond', color='red'),
            name=f"{ticker} Max",
            showlegend=False,
            hovertemplate=f"{ticker}<br>Date=%{{x}}<br>Max=%{{y:.6f}}<extra></extra>"
        ))

        fig_obj.add_trace(go.Scatter(
            x=[min_row['Date']],
            y=[min_row['Value']],
            mode='markers+text',
            text=['Min'],
            textposition='bottom center',
            marker=dict(size=9, symbol='diamond', color='blue'),
            name=f"{ticker} Min",
            showlegend=False,
            hovertemplate=f"{ticker}<br>Date=%{{x}}<br>Min=%{{y:.6f}}<extra></extra>"
        ))

    return fig_obj


def compute_descriptive_statistics(returns_df, cumulative_return_df, trading_days=252):
    rows = []

    for ticker in returns_df.columns:
        s = returns_df[ticker].dropna()
        c = cumulative_return_df[ticker].dropna() if ticker in cumulative_return_df.columns else pd.Series(dtype=float)

        if len(s) == 0:
            rows.append({
                'Ticker': ticker,
                'Mean': np.nan,
                'Std': np.nan,
                'Skewness': np.nan,
                'Kurtosis': np.nan,
                'Annualized Sharpe': np.nan,
                'Cumulative Return': np.nan,
                'Max Drawdown': np.nan
            })
            continue

        mean_val = s.mean()
        std_val = s.std(ddof=1)
        skew_val = s.skew()
        kurt_val = s.kurt()
        sharpe = (mean_val / std_val) * np.sqrt(trading_days) if pd.notna(std_val) and std_val != 0 else np.nan

        if len(c) > 0:
            wealth = 1 + c
            running_max = wealth.cummax()
            drawdown = wealth / running_max - 1
            max_drawdown = drawdown.min()
            cumulative_ret = c.iloc[-1]
        else:
            cumulative_ret = np.nan
            max_drawdown = np.nan

        rows.append({
            'Ticker': ticker,
            'Mean': mean_val,
            'Std': std_val,
            'Skewness': skew_val,
            'Kurtosis': kurt_val,
            'Annualized Sharpe': sharpe,
            'Cumulative Return': cumulative_ret,
            'Max Drawdown': max_drawdown
        })

    return pd.DataFrame(rows)


def compute_financial_metrics(price_wide, returns_df=None, benchmark_series=None, trading_days=252, risk_free_rate_annual=0.0):
    if price_wide is None or price_wide.empty:
        return (
            pd.DataFrame(), pd.DataFrame(), pd.DataFrame(),
            pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
        )

    price_wide = price_wide.sort_index().copy()
    price_wide = price_wide.apply(pd.to_numeric, errors='coerce')

    if returns_df is None or returns_df.empty:
        returns_df_local = price_wide.pct_change()
    else:
        returns_df_local = returns_df.copy().sort_index()
        returns_df_local = returns_df_local.apply(pd.to_numeric, errors='coerce')

    log_returns_df_local = np.log1p(returns_df_local)
    cumulative_return_df_local = (1 + returns_df_local.fillna(0)).cumprod() - 1
    rolling_vol_20_df_local = returns_df_local.rolling(window=20, min_periods=20).std(ddof=1) * np.sqrt(trading_days)

    summary_rows = []

    for col in price_wide.columns:
        s_price = pd.to_numeric(price_wide[col], errors='coerce').dropna()
        s_ret = pd.to_numeric(returns_df_local[col], errors='coerce').dropna() if col in returns_df_local.columns else pd.Series(dtype=float)
        s_log = pd.to_numeric(log_returns_df_local[col], errors='coerce').dropna() if col in log_returns_df_local.columns else pd.Series(dtype=float)
        s_cum = pd.to_numeric(cumulative_return_df_local[col], errors='coerce').dropna() if col in cumulative_return_df_local.columns else pd.Series(dtype=float)

        if len(s_ret) >= 1 and len(s_price) >= 2:
            total_return = (s_price.iloc[-1] / s_price.iloc[0]) - 1 if s_price.iloc[0] != 0 else np.nan
            n_periods = len(s_ret)

            if pd.notna(total_return) and (1 + total_return) > 0 and n_periods > 0:
                annualized_return = (1 + total_return) ** (trading_days / n_periods) - 1
            else:
                annualized_return = np.nan

            daily_vol = s_ret.std(ddof=1) if len(s_ret) > 1 else np.nan
            annualized_vol = daily_vol * np.sqrt(trading_days) if pd.notna(daily_vol) else np.nan
            mean_ret = s_ret.mean()

            var_hist_95 = -np.percentile(s_ret, 5) if len(s_ret) > 0 else np.nan
            var_hist_99 = -np.percentile(s_ret, 1) if len(s_ret) > 0 else np.nan

            mu = s_ret.mean() if len(s_ret) > 0 else np.nan
            sigma = s_ret.std(ddof=1) if len(s_ret) > 1 else np.nan
            z95 = z_score_from_conf(0.95)
            z99 = z_score_from_conf(0.99)

            var_param_95 = -(mu - z95 * sigma) if pd.notna(mu) and pd.notna(sigma) else np.nan
            var_param_99 = -(mu - z99 * sigma) if pd.notna(mu) and pd.notna(sigma) else np.nan

            p5 = np.percentile(s_ret, 5) if len(s_ret) > 0 else np.nan
            p1 = np.percentile(s_ret, 1) if len(s_ret) > 0 else np.nan

            tail_95 = s_ret[s_ret <= p5] if len(s_ret) > 0 else pd.Series(dtype=float)
            tail_99 = s_ret[s_ret <= p1] if len(s_ret) > 0 else pd.Series(dtype=float)

            cvar_hist_95 = -tail_95.mean() if len(tail_95) > 0 else np.nan
            cvar_hist_99 = -tail_99.mean() if len(tail_99) > 0 else np.nan

            max_drawdown = compute_max_drawdown(s_cum)

            excess_return = annualized_return
            sharpe_ratio = (
                (annualized_return - risk_free_rate_annual) / annualized_vol
                if pd.notna(annualized_return) and pd.notna(annualized_vol) and annualized_vol != 0
                else np.nan
            )

            calmar_ratio = (
                annualized_return / abs(max_drawdown)
                if pd.notna(annualized_return) and pd.notna(max_drawdown) and max_drawdown != 0
                else np.nan
            )

            beta_val = np.nan
            treynor_ratio = np.nan
            information_ratio = np.nan
            tracking_error = np.nan

            if benchmark_series is not None:
                aligned_benchmark = benchmark_series.reindex(returns_df_local.index)
                merged = pd.concat([returns_df_local[col], aligned_benchmark], axis=1).dropna()
                if merged.shape[0] > 1:
                    asset_ret = merged.iloc[:, 0]
                    bench_ret = merged.iloc[:, 1]

                    active_ret = asset_ret - bench_ret
                    annualized_active_return = active_ret.mean() * trading_days
                    tracking_error = compute_tracking_error(active_ret, trading_days=trading_days)
                    information_ratio = (
                        annualized_active_return / tracking_error
                        if pd.notna(tracking_error) and tracking_error != 0 else np.nan
                    )

                    beta_val = compute_beta(asset_ret, bench_ret)
                    treynor_ratio = (
                        (annualized_return - risk_free_rate_annual) / beta_val
                        if pd.notna(beta_val) and beta_val != 0 and pd.notna(annualized_return)
                        else np.nan
                    )

                    excess_return = annualized_active_return

            summary_rows.append({
                'Ticker': col,
                'Observations': len(s_ret),
                'Total Return': total_return,
                'Mean Daily Return': mean_ret,
                'Mean Log Return': s_log.mean() if len(s_log) > 0 else np.nan,
                'Annualized Return': annualized_return,
                'Daily Volatility': daily_vol,
                'Annualized Volatility': annualized_vol,
                'VaR 95% (Historical)': var_hist_95,
                'VaR 99% (Historical)': var_hist_99,
                'VaR 95% (Parametric)': var_param_95,
                'VaR 99% (Parametric)': var_param_99,
                'CVaR 95% (Historical)': cvar_hist_95,
                'CVaR 99% (Historical)': cvar_hist_99,
                'Maximum Drawdown': max_drawdown,
                'Sharpe Ratio': sharpe_ratio,
                'Calmar Ratio': calmar_ratio,
                'Beta': beta_val,
                'Treynor Ratio': treynor_ratio,
                'Tracking Error': tracking_error,
                'Information Ratio': information_ratio,
                'Excess Return': excess_return
            })
        else:
            summary_rows.append({
                'Ticker': col,
                'Observations': 0,
                'Total Return': np.nan,
                'Mean Daily Return': np.nan,
                'Mean Log Return': np.nan,
                'Annualized Return': np.nan,
                'Daily Volatility': np.nan,
                'Annualized Volatility': np.nan,
                'VaR 95% (Historical)': np.nan,
                'VaR 99% (Historical)': np.nan,
                'VaR 95% (Parametric)': np.nan,
                'VaR 99% (Parametric)': np.nan,
                'CVaR 95% (Historical)': np.nan,
                'CVaR 99% (Historical)': np.nan,
                'Maximum Drawdown': np.nan,
                'Sharpe Ratio': np.nan,
                'Calmar Ratio': np.nan,
                'Beta': np.nan,
                'Treynor Ratio': np.nan,
                'Tracking Error': np.nan,
                'Information Ratio': np.nan,
                'Excess Return': np.nan
            })

    summary_df_local = pd.DataFrame(summary_rows)
    correlation_matrix_local = returns_df_local.corr()

    return (
        returns_df_local,
        log_returns_df_local,
        cumulative_return_df_local,
        summary_df_local,
        correlation_matrix_local,
        rolling_vol_20_df_local
    )


def compute_rolling_metrics(returns_df, window=30, trading_days=252):
    if returns_df is None or returns_df.empty:
        return pd.DataFrame(), pd.DataFrame()

    returns_df = returns_df.sort_index()
    rolling_mean_return_df = returns_df.rolling(window=window, min_periods=window).mean()
    rolling_volatility_df = returns_df.rolling(window=window, min_periods=window).std(ddof=1) * np.sqrt(trading_days)

    return rolling_mean_return_df, rolling_volatility_df


def format_worksheet(ws):
    header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
    body_font = Font(name='Arial', size=10, color="000000")
    header_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )
    center_alignment = Alignment(horizontal='center', vertical='center')

    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if cell.value is not None and len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except Exception:
                pass
        adjusted_width = min((max_length + 2) * 1.2, 50)
        ws.column_dimensions[column].width = adjusted_width

    if ws.max_row >= 2:
        for cell in ws[2]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_alignment
            cell.border = thin_border

    for row in ws.iter_rows(min_row=3):
        for cell in row:
            cell.font = body_font
            cell.border = thin_border
            if row[0].row % 2 == 0:
                cell.fill = PatternFill(start_color="F7F7F7", end_color="F7F7F7", fill_type="solid")

# =========================================================
# Helper: External Data APIs (World Bank & SEC)
# =========================================================
@st.cache_data(ttl=86400)
def fetch_world_bank_macro():
    indicators = {'NY.GDP.MKTP.KD.ZG': 'GDP Growth (%)', 'FP.CPI.TOTL.ZG': 'Inflation CPI (%)', 'SL.UEM.TOTL.ZS': 'Unemployment (%)'}
    countries = {'US': 'United States', 'CN': 'China', 'JP': 'Japan', 'EU': 'European Union'}
    data = []
    for icode, iname in indicators.items():
        for ccode, cname in countries.items():
            try:
                url = f"http://api.worldbank.org/v2/country/{ccode}/indicator/{icode}?format=json&per_page=5"
                res = requests.get(url, timeout=5).json()
                if len(res) == 2:
                    for item in res[1]:
                        if item['value'] is not None:
                            data.append({'Indicator': iname, 'Country': cname, 'Year': item['date'], 'Value': item['value']})
            except: pass
    return pd.DataFrame(data)

@st.cache_data(ttl=86400)
def get_sec_cik_mapping():
    try:
        headers = {"User-Agent": "ResearchApp/1.0 (info@example.com)"}
        res = requests.get("https://www.sec.gov/files/company_tickers.json", headers=headers, timeout=5).json()
        return {v['ticker']: str(v['cik_str']).zfill(10) for k, v in res.items()}
    except: return {}

def fetch_sec_filings(ticker, cik_map, limit=8):
    """Fetch an expanded array of recent SEC filings including 8-K, Proxies, and Prospectuses."""
    if ticker not in cik_map: return []
    cik = cik_map[ticker]
    try:
        headers = {"User-Agent": "ResearchApp/1.0 (info@example.com)"}
        url = f"https://data.sec.gov/submissions/CIK{cik}.json"
        res = requests.get(url, headers=headers, timeout=5).json()
        recent = res.get('filings', {}).get('recent', {})
        docs = []
        
        target_forms = ['10-K', '10-Q', '8-K', 'DEF 14A', 'DEFA14A', 'SD', '424B2', 'PX14A6G', 'PX14A6N', '1-A POS', 'FWP']
        
        for f, d, a in zip(recent.get('form', []), recent.get('filingDate', []), recent.get('accessionNumber', [])):
            if f in target_forms or f.startswith('424'):
                link = f"https://www.sec.gov/Archives/edgar/data/{int(cik)}/{a.replace('-','')}/{a}-index.html"
                
                # Descriptive labels for professional display
                desc = "Periodic Report"
                if f in ['10-K']: desc = "Annual Report"
                elif f in ['10-Q']: desc = "Quarterly Report"
                elif f in ['8-K']: desc = "Current Report / Material Event"
                elif '14A' in f or '14A6' in f: desc = "Proxy Statement / Shareholder Material"
                elif f == 'SD': desc = "Specialized Disclosure (e.g., Conflict Minerals)"
                elif f.startswith('424') or f in ['FWP', '1-A POS']: desc = "Prospectus / Registration (e.g., Preferred Stock/Debt offerings)"
                
                docs.append({'Form': f, 'Date': d, 'Accession': a, 'Link': link, 'Description': desc})
                if len(docs) >= limit: break
        return docs
    except: return []

def fetch_sec_text_insights(sec_links_list):
    import html
    insights = {
        'Business Overview': "Data not found or could not be parsed.",
        'Risk Factors': "Data not found or could not be parsed.",
        'MD&A': "Data not found or could not be parsed.",
        'CEO Quotes': "No recent quotes found.",
        'Forward Looking': "No forward-looking guidance found."
    }
    
    headers = {"User-Agent": "ResearchApp/1.0 (info@example.com)"}

    def _extract_clean_text(raw_sec_text):
        # 1. Remove binary/attachment documents from EDGAR .txt
        parts = raw_sec_text.split('<DOCUMENT>')
        clean_parts = [parts[0]]
        for p in parts[1:]:
            if re.search(r'<TYPE>(GRAPHIC|ZIP|PDF|EXCEL|XML|JSON)', p, re.IGNORECASE):
                continue
            clean_parts.append(p)
        t = " ".join(clean_parts)
        
        # 2. Remove inline scripts and styles
        t = re.sub(r'<script[^>]*>.*?</script>', ' ', t, flags=re.IGNORECASE | re.DOTALL)
        t = re.sub(r'<style[^>]*>.*?</style>', ' ', t, flags=re.IGNORECASE | re.DOTALL)
        
        # 3. Strip HTML tags
        t = re.sub(r'<[^>]+>', ' ', t)
        
        # 4. Unescape HTML entities
        t = html.unescape(t)
        
        # 5. Normalize whitespace
        t = re.sub(r'\s+', ' ', t)
        
        return t.strip()

    def _clean_snippet(text, length):
        snippet = text[:length].strip()
        # Remove long sequences of symbols (like -------- or _______)
        snippet = re.sub(r'([^\w\s])\1{4,}', ' ', snippet)
        return snippet + "..."

    # --- 1. Process 10-K ---
    ten_k = next((f for f in sec_links_list if f['Form'] == '10-K'), None)
    if ten_k:
        try:
            txt_url = ten_k['Link'].replace('-index.html', '.txt')
            res = requests.get(txt_url, headers=headers, timeout=10)
            raw_text = _extract_clean_text(res.text)
            
            # MD&A
            mda_matches = list(re.finditer(r'ITEM\s+7[\.\:\-\s]*MANAGEMENT\S?S\s+DISCUSSION', raw_text, re.IGNORECASE))
            if not mda_matches:
                mda_matches = list(re.finditer(r'MANAGEMENT\S?S\s+DISCUSSION\s+AND\s+ANALYSIS', raw_text, re.IGNORECASE))
            if mda_matches:
                idx = mda_matches[-1].end()
                insights['MD&A'] = _clean_snippet(raw_text[idx:], 1500)
            
            # Risk Factors
            rf_matches = list(re.finditer(r'ITEM\s+1A[\.\:\-\s]*RISK\s+FACTORS', raw_text, re.IGNORECASE))
            if not rf_matches:
                rf_matches = list(re.finditer(r'RISK\s+FACTORS', raw_text, re.IGNORECASE))
            if rf_matches:
                idx = rf_matches[-1].end()
                insights['Risk Factors'] = _clean_snippet(raw_text[idx:], 1000)
                
            # Business Overview
            bus_matches = list(re.finditer(r'ITEM\s+1[\.\:\-\s]*BUSINESS', raw_text, re.IGNORECASE))
            if bus_matches:
                idx = bus_matches[-1].end()
                insights['Business Overview'] = _clean_snippet(raw_text[idx:], 800)
        except Exception:
            pass

    # --- 2. Process 8-K or 10-Q for CEO Quotes ---
    recent_report = next((f for f in sec_links_list if f['Form'] in ['8-K', '10-Q']), None)
    if recent_report:
        try:
            txt_url = recent_report['Link'].replace('-index.html', '.txt')
            res = requests.get(txt_url, headers=headers, timeout=10)
            raw_text = _extract_clean_text(res.text)
            
            # CEO Quotes Heuristic
            # Look for quote marks with 80-400 characters inside
            quotes = re.findall(r'[“"”]([^”"”]{80,400})[“"”]', raw_text)
            valid_quotes = []
            for q_text in quotes:
                if not q_text or len(q_text.strip()) < 80: continue
                # Heuristic to skip base64 garbage if any slipped through:
                if sum(c.isalpha() for c in q_text) / len(q_text) < 0.6: continue
                if q_text.count(' ') < 5: continue
                
                idx = raw_text.find(q_text)
                context = raw_text[max(0, idx-80) : min(len(raw_text), idx+len(q_text)+80)].lower()
                
                if any(kw in context for kw in ['said', 'stated', 'ceo', 'announced', 'commented', 'president']):
                    valid_quotes.append(q_text.strip())
                    
            if valid_quotes:
                insights['CEO Quotes'] = f'"{valid_quotes[0]}"'
            elif quotes:
                # Fallback to the first decent looking quote
                for q in quotes:
                    if sum(c.isalpha() for c in q) / len(q) > 0.6 and q.count(' ') > 5:
                        insights['CEO Quotes'] = f'"{q.strip()}"'
                        break
                
            # Forward-looking Guidance
            sentences = re.split(r'(?<=[.!?])\s+', raw_text)
            fw_sentences = [s for s in sentences if re.search(r'\b(guidance|outlook|expects|anticipates|projected)\b', s, re.IGNORECASE) and len(s) > 40]
            if fw_sentences:
                insights['Forward Looking'] = " ".join(fw_sentences[:2]).strip() + "."
        except Exception:
            pass

    return insights

@st.cache_data(ttl=86400)
def fetch_sec_company_facts(ticker, cik_map):
    """Fetch SEC EDGAR XBRL company facts to build financial statements."""
    if ticker not in cik_map: return {}
    cik = cik_map[ticker]
    headers = {"User-Agent": "ResearchApp/1.0 (info@example.com)"}
    url = f"https://data.sec.gov/api/xbrl/companyfacts/CIK{cik}.json"
    try:
        res = requests.get(url, headers=headers, timeout=10).json()
        return res.get('facts', {}).get('us-gaap', {})
    except:
        return {}

def extract_financial_data(facts):
    """Extract expanded standard US GAAP metrics from SEC facts for the last 3 fiscal years (10-K)."""
    def get_annual_data(tag_list):
        for tag in tag_list:
            if tag in facts:
                units = facts[tag].get('units', {})
                key_to_use = 'USD' if 'USD' in units else ('USD/shares' if 'USD/shares' in units else None)
                if key_to_use:
                    data = units[key_to_use]
                    # Fetch only Annual reports (10-K, 10-K/A)
                    annual_data = [d for d in data if d.get('form') in ['10-K', '10-K/A']]
                    if annual_data:
                        df = pd.DataFrame(annual_data)
                        if 'fy' in df.columns and 'val' in df.columns:
                            # Keep latest filing per fiscal year
                            df = df.dropna(subset=['fy', 'val']).drop_duplicates(subset=['fy'], keep='last')
                            df = df.sort_values('fy', ascending=False).head(3)
                            return dict(zip(df['fy'].astype(int), df['val']))
        return {}

    # Expanded tags to completely eliminate N/A for major companies like AAPL
    tags_revenue = ['Revenues', 'NetSales', 'SalesRevenueNet', 'SalesRevenueGoodsNet', 'RevenueFromContractWithCustomerExcludingAssessedTax']
    tags_gross_profit = ['GrossProfit', 'GrossMargin']
    tags_op_income = ['OperatingIncomeLoss']
    tags_net_income = ['NetIncomeLoss', 'ProfitLoss']
    tags_assets = ['Assets']
    tags_liabilities = ['Liabilities']
    tags_equity = ['StockholdersEquity', 'StockholdersEquityIncludingPortionAttributableToNoncontrollingInterest']
    tags_op_cf = ['NetCashProvidedByUsedInOperatingActivities']
    tags_eps = ['EarningsPerShareBasic', 'EarningsPerShareDiluted']
    tags_rnd = ['ResearchAndDevelopmentExpense', 'ResearchAndDevelopmentExpenseExcludingAcquiredInProcessCost']
    tags_sga = ['SellingGeneralAndAdministrativeExpense']
    tags_debt = ['LongTermDebt', 'LongTermDebtAndCapitalLeaseObligations', 'LongTermDebtNoncurrent']
    tags_dividends = ['DividendsCash', 'DividendsPaid', 'PaymentsOfDividendsCommonStock']
    tags_dps = ['CommonStockDividendsPerShareDeclared', 'CommonStockDividendsPerShareCashPaid']

    return {
        'Total Revenue': get_annual_data(tags_revenue),
        'Gross Profit': get_annual_data(tags_gross_profit),
        'Operating Income': get_annual_data(tags_op_income),
        'Net Income': get_annual_data(tags_net_income),
        'Total Assets': get_annual_data(tags_assets),
        'Total Liabilities': get_annual_data(tags_assets),
        'Stockholders Equity': get_annual_data(tags_equity),
        'Operating Cash Flow': get_annual_data(tags_op_cf),
        'Earnings Per Share (EPS)': get_annual_data(tags_eps),
        'R&D Expense': get_annual_data(tags_rnd),
        'SG&A Expense': get_annual_data(tags_sga),
        'Long-Term Debt': get_annual_data(tags_debt),
        'Cash Dividends': get_annual_data(tags_dividends),
        'Dividends Per Share': get_annual_data(tags_dps)
    }

@st.cache_data(ttl=3600)
def get_latest_price(ticker, _engine):
    """Fetch the single latest available closing price for a given ticker from WRDS."""
    try:
        query = f"""
        SELECT a.date, a.prc
        FROM crsp.dsf a
        JOIN crsp.msenames b ON a.permno = b.permno
        WHERE b.ticker = '{ticker}'
        ORDER BY a.date DESC LIMIT 1
        """
        with _engine.connect() as conn:
            df = pd.read_sql(query, conn)
        if not df.empty:
            return abs(float(df['prc'].iloc[0]))
        return np.nan
    except Exception:
        return np.nan

# =========================================================
# Docx Wall Street Styling Helpers
# =========================================================
def set_wall_street_table_style(table):
    """
    Apply a standard Wall Street / Academic table style to a python-docx table.
    Removes vertical borders, adds thick top & bottom borders, and light horizontal inner borders.
    Ensures text is properly aligned and table layout is clean.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Autofit settings
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)

    # Remove all existing border settings
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)

    # Rebuild Wall Street borders
    tblBorders = OxmlElement('w:tblBorders')

    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '16')  # Thick line
    top.set(qn('w:color'), '000000')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '16')  # Thick line
    bottom.set(qn('w:color'), '000000')

    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')   # Thin horizontal line
    insideH.set(qn('w:color'), 'A0A0A0')  # Gray

    tblBorders.append(top)
    tblBorders.append(bottom)
    tblBorders.append(insideH)
    tblPr.append(tblBorders)

    # Zebra striping for alternating rows and bold first row
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            # Right align (except first column)
            if cell != row.cells[0]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Bold first row and background color
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'F2F2F2')
                cell._tc.get_or_add_tcPr().append(shading)
            # Even row zebra striping
            elif i % 2 == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'FAFAFA')
                cell._tc.get_or_add_tcPr().append(shading)

def add_chapter_heading(doc, text, level=1):
    """Adds a heading that forces a page break before it, completely replacing doc.add_page_break()"""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.page_break_before = True
    return h


# =========================================================
# Sidebar
# =========================================================
st.sidebar.markdown("---")
APP_PAGE = st.sidebar.radio("🧭 Navigation", ["Financial Terminal", "Report Generator"])
st.sidebar.markdown("---")

with st.sidebar.expander("📘 Financial Metric Reference Table", expanded=False):
    st.markdown("""
This reference table provides a lookup for the financial metrics used in this app, including:
- English Name
- Standard formula
- Interpretation and practical ranges
""")

    guide_metric_df = REFERENCE_DF.copy()

    guide_metric_df = guide_metric_df.rename(columns={
        "Metric": "Metric Code",
        "Formula": "Calculation Formula"
    })

    show_cols = [c for c in [
        "Metric Code",
        "English Name",
        "Calculation Formula",
        "Interpretation",
        "Common Reference Range",
        "How to Judge"
    ] if c in guide_metric_df.columns]

    st.dataframe(guide_metric_df[show_cols], use_container_width=True, hide_index=True)


with st.sidebar.expander("Open-source & Historical Reference Notes", expanded=False):
    st.markdown("""
**Important Note**
- These ranges are **practical reference ranges**, not universal absolute standards.
- They are synthesized from **open quantitative finance practice**, **widely used public market analytics conventions**, and **historical observations often cited in equities and portfolio management**.

**Suggested Interpretation**
- Use them for **relative comparison**, not hard pass/fail thresholds.
- Compare:
  1. same asset across periods
  2. different assets within the same period
  3. return together with risk, not separately
""")

st.sidebar.markdown("---")
st.sidebar.title("🔐 WRDS Terminal Login")

with st.sidebar.form("wrds_login"):
    st.info("Please enter your WRDS credentials.")
    username = st.text_input("WRDS Username")
    password = st.text_input("WRDS Password", type="password")
    login_button = st.form_submit_button("Connect Database")

if 'db_engine' not in st.session_state:
    st.session_state.db_engine = None
if 'connection_status' not in st.session_state:
    st.session_state.connection_status = "Not Connected"

if login_button:
    if username and password:
        with st.spinner("Connecting..."):
            try:
                conn_str = f"postgresql://{username}:{password}@wrds-pgdata.wharton.upenn.edu:9737/wrds"
                engine = create_engine(conn_str)
                with engine.connect() as connection:
                    connection.execute(text("SELECT 1"))
                st.session_state.db_engine = engine
                st.session_state.connection_status = "Connected"
                st.sidebar.success("Connected Successfully!")
            except Exception as e:
                st.sidebar.error(f"Connection Failed: {e}")
                st.session_state.db_engine = None
                st.session_state.connection_status = "Not Connected"
    else:
        st.sidebar.warning("Please enter username and password")

st.sidebar.markdown(f"Status: **{st.session_state.connection_status}**")

st.sidebar.markdown("---")
st.sidebar.subheader("📉 Rolling Analysis Settings")
enable_rolling_analysis = st.sidebar.checkbox("Enable Rolling Analysis", value=False)
rolling_window = st.sidebar.selectbox(
    "Rolling Window (Days)",
    options=[20, 30, 60, 90],
    index=1,
    disabled=not enable_rolling_analysis
)
rolling_show_return = st.sidebar.checkbox(
    "Show Rolling Mean Return",
    value=True,
    disabled=not enable_rolling_analysis
)
rolling_show_vol = st.sidebar.checkbox(
    "Show Rolling Volatility",
    value=True,
    disabled=not enable_rolling_analysis
)


# =========================================================
# MODULE 1: Financial Terminal
# =========================================================
if APP_PAGE == "Financial Terminal":

    st.title("📊 WRDS Financial Terminal")
    st.caption("WRDS CRSP Multi-Company Analysis · Rolling Risk · Excel Export · Bloomberg UI Style")

    if st.session_state.db_engine is None:
        st.warning("⚠️ Please connect WRDS in the sidebar first.")
        st.stop()

    engine = st.session_state.db_engine

    st.header("🔍 Query & Filter Configuration")

    db_columns, display_options, display_to_column = get_table_columns(engine)

    with st.container():
        top_filter_col1, top_filter_col2, top_filter_col3 = st.columns([2, 1, 1])

        with top_filter_col1:
            tickers_input = st.text_input(
                "Tickers (comma separated)",
                value="", placeholder="Example: AAPL, MSFT, GOOGL",
                help="Example: AAPL, MSFT, GOOGL"
            ).upper()
            ticker_list = [t.strip() for t in tickers_input.split(",") if t.strip()]

        with top_filter_col2:
            start_date = st.date_input(
                "Start Date",
                value=datetime.date(2023, 1, 1)
            )

        with top_filter_col3:
            end_date = st.date_input(
                "End Date",
                value=datetime.date.today()
            )

    st.markdown("---")

    st.subheader("📋 Raw Data Column Selection")

    default_cols = ['prc', 'vol', 'ret', 'cfacpr']
    default_display = []
    for c in default_cols:
        for opt in display_options:
            if opt.startswith(f"{c} ("):
                default_display.append(opt)
                break

    selected_display = st.multiselect(
        "Select raw columns from WRDS (crsp.dsf)",
        options=display_options,
        default=[],
        help="Choose the raw WRDS columns you want to query and display/export."
    )

    selected_columns = [display_to_column[opt] for opt in selected_display if opt in display_to_column]

    st.markdown("---")

    st.subheader("🧮 Financial Analysis Feature Selection")

    financial_feature_options = [f"{k} ({v})" for k, v in FIN_ANALYSIS_MAPPING.items()]

    default_financial_features = [
        f"daily_return_py ({FIN_ANALYSIS_MAPPING['daily_return_py']})",
        f"log_return_py ({FIN_ANALYSIS_MAPPING['log_return_py']})",
        f"annualized_return_py ({FIN_ANALYSIS_MAPPING['annualized_return_py']})",
        f"cumulative_return_py ({FIN_ANALYSIS_MAPPING['cumulative_return_py']})",
        f"volatility_annualized_py ({FIN_ANALYSIS_MAPPING['volatility_annualized_py']})",
        f"var_hist_95_py ({FIN_ANALYSIS_MAPPING['var_hist_95_py']})",
        f"cvar_hist_95_py ({FIN_ANALYSIS_MAPPING['cvar_hist_95_py']})",
        f"max_drawdown_py ({FIN_ANALYSIS_MAPPING['max_drawdown_py']})",
        f"sharpe_ratio_py ({FIN_ANALYSIS_MAPPING['sharpe_ratio_py']})",
        f"calmar_ratio_py ({FIN_ANALYSIS_MAPPING['calmar_ratio_py']})",
        f"beta_py ({FIN_ANALYSIS_MAPPING['beta_py']})",
        f"treynor_ratio_py ({FIN_ANALYSIS_MAPPING['treynor_ratio_py']})",
        f"information_ratio_py ({FIN_ANALYSIS_MAPPING['information_ratio_py']})"
    ]

    selected_financial_features_display = st.multiselect(
        "Select calculated financial analysis items",
        options=financial_feature_options,
        default=[],
        help="Choose which calculated metrics should be computed, displayed, and optionally exported."
    )

    selected_financial_features = [x.split(" (")[0] for x in selected_financial_features_display]

    analysis_col1, analysis_col2, analysis_col3, analysis_col4 = st.columns(4)
    with analysis_col1:
        trading_days_per_year = st.number_input(
            "Trading Days per Year",
            min_value=200,
            max_value=366,
            value=252,
            step=1
        )
    with analysis_col2:
        risk_confidence_level = st.selectbox(
            "Primary Risk Confidence Level",
            options=[0.90, 0.95, 0.99],
            index=1
        )
    with analysis_col3:
        benchmark_mode = st.selectbox(
            "Excess Return Benchmark",
            options=["Zero Rate / 0%", "Cross-sectional Mean Return"]
        )
    with analysis_col4:
        risk_free_rate_annual = st.number_input(
            "Annual Risk-free Rate",
            min_value=-0.20,
            max_value=0.30,
            value=0.02,
            step=0.005,
            format="%.3f"
        )

    st.info("📘 Open the sidebar Metric Guide to view explanations, formulas, and practical reference ranges.")

    st.markdown("---")

    st.subheader("📊 Visualization Configuration")

    viz_col1, viz_col2, viz_col3 = st.columns([2, 1, 1])

    with viz_col1:
        comparison_metric = st.selectbox(
            "Select comparison metric",
            ["Price", "Daily Return", "Cumulative Return", "Rolling Volatility (20D, Annualized)"]
        )

    with viz_col2:
        comparison_chart_type = st.radio(
            "Chart Type",
            ["Line", "Bar"],
            horizontal=True
        )

    with viz_col3:
        show_data_preview = st.checkbox("Show Data Preview Tables", value=True)

    st.markdown("---")

    st.subheader("🧹 Data Cleaning & Price Adjustment")

    clean_col1, clean_col2, clean_col3 = st.columns(3)

    with clean_col1:
        missing_value_method = st.selectbox(
            "Missing/Invalid Data Handling",
            ["Forward Fill", "Mean Fill", "Drop Missing"],
            index=0
        )

    with clean_col2:
        use_adjusted_price = st.checkbox("Use Adjusted Price", value=True)

    with clean_col3:
        adjustment_mode = st.selectbox(
            "Adjustment Formula",
            ["multiply", "divide"],
            index=0
        )

    st.markdown("---")

    st.subheader("⚙️ Export Settings")

    export_col1, export_col2 = st.columns([2, 1])

    BASE_EXPORT_OPTIONS = [
        "Per-Ticker Sheets",
        "Summary Comparison",
        "Raw WRDS Data",
        "Enhanced Long Format",
        "Financial Time Series",
        "Financial Summary",
        "Rolling Time Series",
        "Rolling Summary",
        "Descriptive Stats",
        "Data Issues",
        "Adjusted Prices",
        "Correlation Matrix",
        "Metric Guide",
        "Charts"
    ]

    def build_dynamic_export_options(
        selected_columns,
        selected_financial_features,
        enable_rolling_analysis,
        include_chart_in_excel
    ):
        export_options = []
        export_options.extend(["Per-Ticker Sheets", "Summary Comparison", "Metric Guide"])

        if selected_columns:
            export_options.extend(["Raw WRDS Data", "Enhanced Long Format", "Descriptive Stats", "Data Issues", "Adjusted Prices", "Correlation Matrix"])

        if selected_financial_features:
            export_options.extend(["Financial Time Series", "Financial Summary"])

        if enable_rolling_analysis:
            export_options.extend(["Rolling Time Series", "Rolling Summary"])

        if include_chart_in_excel:
            export_options.append("Charts")

        export_options.extend(BASE_EXPORT_OPTIONS)
        seen = set()
        export_options = [x for x in export_options if not (x in seen or seen.add(x))]
        return export_options


    def build_default_export_selection(
        export_options,
        selected_columns,
        selected_financial_features,
        enable_rolling_analysis,
        include_chart_in_excel
    ):
        defaults = ["Metric Guide", "Summary Comparison"]
        if selected_columns: defaults += ["Raw WRDS Data", "Adjusted Prices", "Correlation Matrix"]
        if selected_financial_features: defaults += ["Financial Summary", "Financial Time Series"]
        if enable_rolling_analysis: defaults += ["Rolling Summary", "Rolling Time Series"]
        if include_chart_in_excel: defaults += ["Charts"]

        defaults = [x for x in defaults if x in export_options]
        if not defaults: defaults = ["Metric Guide"]
        return defaults


    with export_col2:
        include_chart_in_excel = st.checkbox(
            "Export Charts to Excel",
            value=True,
            help="If enabled, charts will be embedded into the Excel report (requires kaleido)."
        )

        export_file_prefix = st.text_input(
            "Export File Name",
            value="WRDS_Report",
            help="The exported Excel file name prefix."
        )

    with export_col1:
        st.markdown("**Export options automatically match your selected analysis modules.**")

        export_options_all = build_dynamic_export_options(
            selected_columns=selected_columns,
            selected_financial_features=selected_financial_features,
            enable_rolling_analysis=enable_rolling_analysis,
            include_chart_in_excel=include_chart_in_excel
        )

        default_export_sections = build_default_export_selection(
            export_options=export_options_all,
            selected_columns=selected_columns,
            selected_financial_features=selected_financial_features,
            enable_rolling_analysis=enable_rolling_analysis,
            include_chart_in_excel=include_chart_in_excel
        )

        selected_export_sections = st.multiselect(
            "Export Content",
            options=export_options_all,
            default=default_export_sections,
            help="Only the selected sections will be exported into the Excel file. You may freely add/remove items."
        )

    st.markdown("---")

    run_col1, run_col2 = st.columns([1, 5])
    with run_col1:
        run_analysis = st.button("🚀 Start Analysis", use_container_width=True)
    with run_col2:
        st.empty()

    if run_analysis:
        if not ticker_list:
            st.warning("Please enter at least one ticker.")
            st.stop()

        if (not selected_columns and not selected_financial_features and not enable_rolling_analysis) and (not selected_export_sections):
            st.warning("Please select at least one raw column / financial feature / rolling analysis OR select export content.")
            st.stop()

        if start_date > end_date:
            st.warning("Start Date cannot be later than End Date.")
            st.stop()

        def map_ticker_by_date(row, names_df):
            sub = names_df[
                (names_df['permno'] == row['permno']) &
                (names_df['namedt'] <= row['date']) &
                (names_df['nameendt'] >= row['date'])
            ]
            if not sub.empty:
                return sub.iloc[0]['ticker']
            sub2 = names_df[names_df['permno'] == row['permno']]
            if not sub2.empty:
                return sub2.iloc[0]['ticker']
            return np.nan

        try:
            with st.spinner(f"Querying data for {', '.join(ticker_list)} ..."):
                required_cols = ['date', 'permno', 'prc', 'ret']

                if 'vol' in selected_columns and 'vol' not in required_cols:
                    required_cols.append('vol')

                if use_adjusted_price and 'cfacpr' not in required_cols:
                    required_cols.append('cfacpr')

                query_cols = list(dict.fromkeys(selected_columns + required_cols))
                ticker_str = "', '".join(ticker_list)

                permno_lookup_sql = f"""
                SELECT DISTINCT permno, ticker, namedt, nameendt
                FROM crsp.msenames
                WHERE ticker IN ('{ticker_str}')
                  AND namedt <= '{end_date}'
                  AND nameendt >= '{start_date}'
                """

                with engine.connect() as conn:
                    lookup_df = pd.read_sql(permno_lookup_sql, conn)

                if lookup_df.empty:
                    st.warning("No valid PERMNO mapping found for the selected tickers and date range.")
                    st.stop()

                lookup_df['ticker'] = lookup_df['ticker'].astype(str).str.upper().str.strip()
                lookup_df['namedt'] = pd.to_datetime(lookup_df['namedt'], errors='coerce')
                lookup_df['nameendt'] = pd.to_datetime(lookup_df['nameendt'], errors='coerce')

                permno_list = sorted(lookup_df['permno'].dropna().unique().tolist())
                if not permno_list:
                    st.warning("No valid PERMNO values found.")
                    st.stop()

                permno_str = ", ".join([str(int(x)) for x in permno_list])

                sql = f"""
                SELECT {', '.join(query_cols)}
                FROM crsp.dsf
                WHERE permno IN ({permno_str})
                  AND date >= '{start_date}'
                  AND date <= '{end_date}'
                ORDER BY permno, date
                """

                with engine.connect() as conn:
                    df_raw = pd.read_sql(sql, conn)

                if df_raw.empty:
                    st.warning("No data found in crsp.dsf for the selected tickers and date range.")
                    st.stop()

                df_raw['date'] = pd.to_datetime(df_raw['date'], errors='coerce')
                df_raw['permno'] = pd.to_numeric(df_raw['permno'], errors='coerce')

                if 'prc' in df_raw.columns:
                    df_raw['prc'] = pd.to_numeric(df_raw['prc'], errors='coerce').abs()

                if 'vol' in df_raw.columns:
                    df_raw['vol'] = pd.to_numeric(df_raw['vol'], errors='coerce')

                if 'cfacpr' in df_raw.columns:
                    df_raw['cfacpr'] = pd.to_numeric(df_raw['cfacpr'], errors='coerce')

                if 'ret' in df_raw.columns:
                    df_raw['ret'] = pd.to_numeric(df_raw['ret'], errors='coerce')

                df_raw = df_raw.dropna(subset=['date', 'permno']).copy()

                if df_raw.empty:
                    st.warning("No valid rows remain after parsing date/permno.")
                    st.stop()

                df_raw['TICKER'] = df_raw.apply(lambda r: map_ticker_by_date(r, lookup_df), axis=1)
                df_raw = df_raw.dropna(subset=['TICKER']).copy()
                df_raw['TICKER'] = df_raw['TICKER'].astype(str).str.upper().str.strip()

                if df_raw.empty:
                    st.warning("Data was retrieved, but no valid ticker mapping could be established.")
                    st.stop()

                original_missing_dates_df = detect_missing_dates(
                    df_raw,
                    date_col='date',
                    ticker_col='TICKER'
                )

                df_raw, invalid_issue_df = clean_missing_and_anomalies(
                    df_raw,
                    price_col='prc',
                    volume_col='vol' if 'vol' in df_raw.columns else None,
                    return_col='ret' if 'ret' in df_raw.columns else None,
                    method=missing_value_method
                )

                if df_raw.empty:
                    st.error("No valid data remains after cleaning.")
                    st.stop()

                if use_adjusted_price:
                    df_raw = apply_price_adjustment(
                        df_raw,
                        price_col='prc',
                        factor_col='cfacpr',
                        mode=adjustment_mode
                    )
                else:
                    df_raw['adj_prc'] = df_raw['prc']

                if not original_missing_dates_df.empty or not invalid_issue_df.empty:
                    data_issues_df = pd.concat(
                        [original_missing_dates_df, invalid_issue_df],
                        axis=0,
                        ignore_index=True
                    )
                else:
                    data_issues_df = pd.DataFrame(columns=['TICKER', 'date', 'Issue'])

                if not data_issues_df.empty:
                    keep_issue_cols = [c for c in ['TICKER', 'date', 'Issue'] if c in data_issues_df.columns]
                    data_issues_df = data_issues_df[keep_issue_cols].drop_duplicates().sort_values(keep_issue_cols[:2])

                grouped_data = df_raw.groupby('TICKER')

                raw_data_cols = [c for c in selected_columns if c not in ['date', 'permno']]
                if 'adj_prc' not in raw_data_cols and 'adj_prc' in df_raw.columns:
                    raw_data_cols.append('adj_prc')

                if raw_data_cols:
                    df_indexed = df_raw.set_index(['date', 'TICKER'])
                    existing_cols = [c for c in raw_data_cols if c in df_indexed.columns]
                    df_pivot = df_indexed[existing_cols].unstack('TICKER')
                    df_pivot.columns = [f"{col}_{ticker}" for col, ticker in df_pivot.columns]
                    df_pivot.index.name = 'Date'
                    df_pivot = df_pivot.reset_index()
                else:
                    df_pivot = pd.DataFrame({'Date': sorted(df_raw['date'].dropna().unique())})

                price_wide = None
                price_value_col = 'adj_prc' if 'adj_prc' in df_raw.columns else 'prc'

                if price_value_col in df_raw.columns:
                    price_wide = df_raw.pivot_table(
                        index='date',
                        columns='TICKER',
                        values=price_value_col,
                        aggfunc='last'
                    ).sort_index()

                returns_from_wrds = pd.DataFrame()
                if 'ret' in df_raw.columns:
                    ret_long = df_raw[['date', 'TICKER', 'ret']].copy()
                    ret_long['ret'] = pd.to_numeric(ret_long['ret'], errors='coerce')
                    ret_long = ret_long.dropna(subset=['date', 'TICKER'])
                    returns_from_wrds = ret_long.pivot_table(
                        index='date',
                        columns='TICKER',
                        values='ret',
                        aggfunc='last'
                    ).sort_index()

                returns_df = pd.DataFrame()
                log_returns_df = pd.DataFrame()
                cumulative_return_df = pd.DataFrame()
                rolling_vol_20_df = pd.DataFrame()
                summary_df = pd.DataFrame()
                correlation_matrix = pd.DataFrame()
                financial_timeseries_export = pd.DataFrame()
                enhanced_long_df = pd.DataFrame()
                descriptive_stats_df = pd.DataFrame()
                adjusted_price_export_df = pd.DataFrame()

                rolling_mean_return_df = pd.DataFrame()
                rolling_volatility_df = pd.DataFrame()
                rolling_mean_long_df = pd.DataFrame()
                rolling_vol_long_df = pd.DataFrame()
                rolling_export_df = pd.DataFrame()
                rolling_summary_df = pd.DataFrame()

                benchmark_series = None
                if benchmark_mode == "Cross-sectional Mean Return":
                    if returns_from_wrds is not None and not returns_from_wrds.empty:
                        benchmark_series = returns_from_wrds.mean(axis=1)
                    elif price_wide is not None and not price_wide.empty:
                        benchmark_series = price_wide.pct_change().mean(axis=1)

                if len(selected_financial_features) > 0 or comparison_metric in [
                    "Price", "Daily Return", "Cumulative Return", "Rolling Volatility (20D, Annualized)"
                ] or enable_rolling_analysis:

                    if price_wide is None or price_wide.empty:
                        st.error("Financial analysis or rolling analysis requires valid price data, but it is unavailable.")
                        st.stop()

                    returns_df, log_returns_df, cumulative_return_df, summary_df, correlation_matrix, rolling_vol_20_df = compute_financial_metrics(
                        price_wide=price_wide,
                        returns_df=returns_from_wrds if not returns_from_wrds.empty else None,
                        benchmark_series=benchmark_series,
                        trading_days=trading_days_per_year,
                        risk_free_rate_annual=risk_free_rate_annual
                    )

                    financial_timeseries_export = pd.DataFrame(index=price_wide.index)

                    if 'daily_return_py' in selected_financial_features and not returns_df.empty:
                        tmp = returns_df.copy()
                        tmp.columns = [f"daily_return_{c}" for c in tmp.columns]
                        financial_timeseries_export = pd.concat([financial_timeseries_export, tmp], axis=1)

                    if 'log_return_py' in selected_financial_features and not log_returns_df.empty:
                        tmp = log_returns_df.copy()
                        tmp.columns = [f"log_return_{c}" for c in tmp.columns]
                        financial_timeseries_export = pd.concat([financial_timeseries_export, tmp], axis=1)

                    if 'cumulative_return_py' in selected_financial_features and not cumulative_return_df.empty:
                        tmp = cumulative_return_df.copy()
                        tmp.columns = [f"cumulative_return_{c}" for c in tmp.columns]
                        financial_timeseries_export = pd.concat([financial_timeseries_export, tmp], axis=1)

                    if 'volatility_annualized_py' in selected_financial_features and not rolling_vol_20_df.empty:
                        tmp = rolling_vol_20_df.copy()
                        tmp.columns = [f"rolling_vol_20d_annualized_{c}" for c in tmp.columns]
                        financial_timeseries_export = pd.concat([financial_timeseries_export, tmp], axis=1)

                    financial_timeseries_export.index.name = 'Date'
                    financial_timeseries_export = financial_timeseries_export.reset_index()

                    if not summary_df.empty:
                        keep_cols = ['Ticker', 'Observations']

                        metric_to_summary_col = {
                            'annualized_return_py': 'Annualized Return',
                            'cumulative_return_py': 'Total Return',
                            'excess_return_py': 'Excess Return',
                            'volatility_daily_py': 'Daily Volatility',
                            'volatility_annualized_py': 'Annualized Volatility',
                            'var_hist_95_py': 'VaR 95% (Historical)',
                            'var_hist_99_py': 'VaR 99% (Historical)',
                            'var_param_95_py': 'VaR 95% (Parametric)',
                            'var_param_99_py': 'VaR 99% (Parametric)',
                            'cvar_hist_95_py': 'CVaR 95% (Historical)',
                            'cvar_hist_99_py': 'CVaR 99% (Historical)',
                            'max_drawdown_py': 'Maximum Drawdown',
                            'sharpe_ratio_py': 'Sharpe Ratio',
                            'calmar_ratio_py': 'Calmar Ratio',
                            'beta_py': 'Beta',
                            'treynor_ratio_py': 'Treynor Ratio',
                            'information_ratio_py': 'Information Ratio'
                        }

                        for feat in selected_financial_features:
                            if feat in metric_to_summary_col:
                                keep_cols.append(metric_to_summary_col[feat])

                        keep_cols = [c for c in list(dict.fromkeys(keep_cols)) if c in summary_df.columns]
                        summary_df = summary_df[keep_cols]

                    enhanced_long_df = build_long_format(
                        comparison_metric,
                        price_wide,
                        returns_df,
                        cumulative_return_df,
                        rolling_vol_20_df
                    )

                    if not returns_df.empty and not cumulative_return_df.empty:
                        descriptive_stats_df = compute_descriptive_statistics(
                            returns_df,
                            cumulative_return_df,
                            trading_days=trading_days_per_year
                        )

                    if price_wide is not None and not price_wide.empty:
                        adjusted_price_export_df = price_wide.copy()
                        adjusted_price_export_df.index.name = 'Date'
                        adjusted_price_export_df = adjusted_price_export_df.reset_index()

                    if enable_rolling_analysis:
                        if returns_df is None or returns_df.empty:
                            st.warning("Rolling analysis requires valid return data, but returns_df is empty.")
                        else:
                            returns_df = returns_df.sort_index()

                            rolling_mean_return_df, rolling_volatility_df = compute_rolling_metrics(
                                returns_df,
                                window=rolling_window,
                                trading_days=trading_days_per_year
                            )

                            rolling_mean_long_df = build_long_format_from_df(rolling_mean_return_df)
                            rolling_vol_long_df = build_long_format_from_df(rolling_volatility_df)

                            rolling_export_df = pd.DataFrame(index=returns_df.index)

                            if rolling_show_return and not rolling_mean_return_df.empty:
                                tmp = rolling_mean_return_df.copy()
                                tmp.columns = [f"rolling_mean_return_{rolling_window}d_{c}" for c in tmp.columns]
                                rolling_export_df = pd.concat([rolling_export_df, tmp], axis=1)

                            if rolling_show_vol and not rolling_volatility_df.empty:
                                tmp = rolling_volatility_df.copy()
                                tmp.columns = [f"rolling_volatility_{rolling_window}d_{c}" for c in tmp.columns]
                                rolling_export_df = pd.concat([rolling_export_df, tmp], axis=1)

                            rolling_export_df.index.name = 'Date'
                            rolling_export_df = rolling_export_df.reset_index()

                            rolling_summary_rows = []
                            for ticker in returns_df.columns:
                                row = {'Ticker': ticker, 'Rolling Window': rolling_window}

                                if rolling_show_return and ticker in rolling_mean_return_df.columns:
                                    s1 = rolling_mean_return_df[ticker].dropna()
                                    row['Rolling Mean Return Avg'] = s1.mean() if not s1.empty else np.nan
                                    row['Rolling Mean Return Max'] = s1.max() if not s1.empty else np.nan
                                    row['Rolling Mean Return Min'] = s1.min() if not s1.empty else np.nan

                                if rolling_show_vol and ticker in rolling_volatility_df.columns:
                                    s2 = rolling_volatility_df[ticker].dropna()
                                    row['Rolling Volatility Avg'] = s2.mean() if not s2.empty else np.nan
                                    row['Rolling Volatility Max'] = s2.max() if not s2.empty else np.nan
                                    row['Rolling Volatility Min'] = s2.min() if not s2.empty else np.nan

                                rolling_summary_rows.append(row)

                            rolling_summary_df = pd.DataFrame(rolling_summary_rows)

            st.success(f"Successfully retrieved {len(df_raw)} records!")

            st.markdown("---")
            st.subheader("📌 Key Performance Indicators (KPI)")

            if not summary_df.empty:
                avg_return = summary_df["Annualized Return"].mean() if "Annualized Return" in summary_df.columns else np.nan
                avg_vol = summary_df["Annualized Volatility"].mean() if "Annualized Volatility" in summary_df.columns else np.nan
                avg_sharpe = summary_df["Sharpe Ratio"].mean() if "Sharpe Ratio" in summary_df.columns else np.nan
                avg_mdd = summary_df["Maximum Drawdown"].mean() if "Maximum Drawdown" in summary_df.columns else np.nan

                k1, k2, k3, k4 = st.columns(4)
                k1.metric("Avg Annual Return", format_percentage(avg_return))
                k2.metric("Avg Annual Volatility", format_percentage(avg_vol))
                k3.metric("Avg Sharpe", format_float(avg_sharpe))
                k4.metric("Avg Max Drawdown", format_percentage(avg_mdd))
            else:
                st.info("KPI data is unavailable because the financial summary table is empty.")

            st.subheader("🧹 Data Quality / Cleaning Result")
            if not data_issues_df.empty:
                st.warning("Detected missing/invalid records before or during cleaning.")
                if show_data_preview:
                    st.dataframe(data_issues_df, use_container_width=True)
            else:
                st.success("No missing or invalid price/volume/return records detected.")

            st.subheader("📈 Raw Price Visualization")
            price_cols = [c for c in df_pivot.columns if 'prc_' in c.lower() or 'adj_prc_' in c.lower()] if not df_pivot.empty else []

            raw_fig = None
            colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#17becf']

            if price_cols and 'Date' in df_pivot.columns:
                raw_fig = go.Figure()
                for idx, col in enumerate(price_cols):
                    ticker_name = col.split('_')[-1]
                    color = colors[idx % len(colors)]
                    raw_fig.add_trace(go.Scatter(
                        x=df_pivot['Date'],
                        y=df_pivot[col],
                        name=ticker_name,
                        mode='lines',
                        line=dict(color=color, width=2)
                    ))

                raw_fig.update_layout(
                    title="Stock Price Comparison",
                    xaxis_title="Date",
                    yaxis_title="Price",
                    template="plotly_dark",
                    hovermode="x unified"
                )
                st.plotly_chart(raw_fig, use_container_width=True)
            else:
                st.info("'prc' / 'adj_prc' price column not available, skipping raw price chart.")

            st.subheader("📊 Enhanced Comparison")
            compare_fig = None

            if not enhanced_long_df.empty:
                if comparison_chart_type == "Line":
                    compare_fig = px.line(
                        enhanced_long_df,
                        x="Date",
                        y="Value",
                        color="Ticker",
                        title=f"{comparison_metric} Comparison Across Companies"
                    )
                else:
                    compare_fig = px.bar(
                        enhanced_long_df,
                        x="Date",
                        y="Value",
                        color="Ticker",
                        barmode="group",
                        title=f"{comparison_metric} Comparison Across Companies"
                    )

                compare_fig = add_max_min_annotations(compare_fig, enhanced_long_df)
                compare_fig.update_layout(
                    xaxis_title="Date",
                    yaxis_title=comparison_metric,
                    hovermode="x unified",
                    template="plotly_dark",
                    legend_title="Company"
                )
                st.plotly_chart(compare_fig, use_container_width=True)

                if show_data_preview:
                    st.markdown("**Long-format comparison data**")
                    st.dataframe(enhanced_long_df, use_container_width=True)
            else:
                st.info("No data available for enhanced comparison visualization.")

            rolling_return_fig = None
            rolling_vol_fig = None

            if enable_rolling_analysis:
                st.subheader("📉 Rolling Analysis")
                st.caption(f"Rolling Window Selected: {rolling_window} Days")

                if returns_df is None or returns_df.empty:
                    st.warning("Rolling analysis could not be displayed because no valid return data is available.")
                else:
                    if show_data_preview:
                        st.markdown("**Underlying Return Data Used for Rolling (from WRDS CRSP dsf.ret)**")
                        st.dataframe(returns_df.reset_index(), use_container_width=True)

                    if rolling_show_return:
                        if not rolling_mean_return_df.empty and not rolling_mean_long_df.empty:
                            st.markdown("**Rolling Mean Return Time Series**")
                            if show_data_preview:
                                st.dataframe(rolling_mean_return_df.reset_index(), use_container_width=True)

                            rolling_return_fig = px.line(
                                rolling_mean_long_df,
                                x="Date",
                                y="Value",
                                color="Ticker",
                                title=f"Rolling Mean Return ({rolling_window}D)"
                            )
                            rolling_return_fig = add_max_min_annotations(rolling_return_fig, rolling_mean_long_df)
                            rolling_return_fig.update_layout(
                                xaxis_title="Date",
                                yaxis_title="Rolling Mean Return",
                                hovermode="x unified",
                                template="plotly_dark"
                            )
                            st.plotly_chart(rolling_return_fig, use_container_width=True)
                        else:
                            st.info(
                                f"Rolling Mean Return has no output yet. "
                                f"This usually means each ticker has fewer than {rolling_window} valid return observations."
                            )

                    if rolling_show_vol:
                        if not rolling_volatility_df.empty and not rolling_vol_long_df.empty:
                            st.markdown("**Rolling Volatility Time Series**")
                            if show_data_preview:
                                st.dataframe(rolling_volatility_df.reset_index(), use_container_width=True)

                            rolling_vol_fig = px.line(
                                rolling_vol_long_df,
                                x="Date",
                                y="Value",
                                color="Ticker",
                                title=f"Rolling Volatility ({rolling_window}D, Annualized)"
                            )
                            rolling_vol_fig = add_max_min_annotations(rolling_vol_fig, rolling_vol_long_df)
                            rolling_vol_fig.update_layout(
                                xaxis_title="Date",
                                yaxis_title="Rolling Volatility",
                                hovermode="x unified",
                                template="plotly_dark"
                            )
                            st.plotly_chart(rolling_vol_fig, use_container_width=True)
                        else:
                            st.info(
                                f"Rolling Volatility has no output yet. "
                                f"This usually means each ticker has fewer than {rolling_window} valid return observations."
                            )

                    if not rolling_summary_df.empty:
                        st.markdown("**Rolling Analysis Summary**")
                        st.dataframe(rolling_summary_df, use_container_width=True)

                    if not rolling_export_df.empty and show_data_preview:
                        st.markdown("**Rolling Analysis Export Preview**")
                        st.dataframe(rolling_export_df, use_container_width=True)

            cum_fig = None
            heatmap_fig = None

            if (len(selected_financial_features) > 0 or comparison_metric in [
                "Daily Return", "Cumulative Return", "Rolling Volatility (20D, Annualized)"
            ]) and price_wide is not None:
                st.subheader("🧮 Financial Analysis Results")

                if 'daily_return_py' in selected_financial_features and not returns_df.empty and show_data_preview:
                    st.markdown("**Daily Return Time Series**")
                    st.dataframe(returns_df.reset_index(), use_container_width=True)

                if 'log_return_py' in selected_financial_features and not log_returns_df.empty and show_data_preview:
                    st.markdown("**Log Return Time Series**")
                    st.dataframe(log_returns_df.reset_index(), use_container_width=True)

                if 'cumulative_return_py' in selected_financial_features and not cumulative_return_df.empty:
                    if show_data_preview:
                        st.markdown("**Cumulative Return Time Series**")
                        st.dataframe(cumulative_return_df.reset_index(), use_container_width=True)

                    cum_fig = go.Figure()
                    for idx, col in enumerate(cumulative_return_df.columns):
                        color = colors[idx % len(colors)]
                        cum_fig.add_trace(go.Scatter(
                            x=cumulative_return_df.index,
                            y=cumulative_return_df[col],
                            name=col,
                            mode='lines',
                            line=dict(color=color, width=2)
                        ))
                    cum_fig.update_layout(
                        title="Cumulative Return Comparison",
                        xaxis_title="Date",
                        yaxis_title="Cumulative Return",
                        template="plotly_dark",
                        hovermode="x unified"
                    )
                    st.plotly_chart(cum_fig, use_container_width=True)

                if not summary_df.empty:
                    st.markdown("**Financial Summary Metrics**")
                    st.dataframe(summary_df, use_container_width=True)

                if not descriptive_stats_df.empty:
                    st.markdown("**Descriptive Statistics Summary**")
                    st.dataframe(descriptive_stats_df, use_container_width=True)

                st.markdown("**Correlation Matrix of Daily Returns**")
                if not correlation_matrix.empty:
                    st.dataframe(correlation_matrix, use_container_width=True)

                    heatmap_fig = px.imshow(
                        correlation_matrix,
                        text_auto=".2f",
                        color_continuous_scale='RdBu_r',
                        zmin=-1,
                        zmax=1,
                        aspect='auto',
                        title='Correlation Heatmap of Daily Returns'
                    )
                    heatmap_fig.update_traces(
                        hovertemplate="Company X: %{x}<br>Company Y: %{y}<br>Correlation: %{z:.4f}<extra></extra>"
                    )
                    heatmap_fig.update_layout(
                        xaxis_title="Company",
                        yaxis_title="Company"
                    )
                    st.plotly_chart(heatmap_fig, use_container_width=True)

            st.subheader("📋 Raw Comparison Data (Pivot Table)")
            if show_data_preview:
                st.dataframe(df_pivot, use_container_width=True)

            st.subheader("📄 Raw WRDS Data (Long Format)")
            if show_data_preview:
                st.dataframe(df_raw, use_container_width=True)

            if not adjusted_price_export_df.empty and show_data_preview:
                st.subheader("🔧 Adjusted Price Data")
                st.dataframe(adjusted_price_export_df, use_container_width=True)

            st.subheader("📥 Export Data")

            output = io.BytesIO()
            temp_files = []

            try:
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    workbook = writer.book

                    if "Per-Ticker Sheets" in selected_export_sections:
                        for ticker, group_df in grouped_data:
                            if pd.isna(ticker):
                                continue

                            cols_to_export = [c for c in selected_columns if c in group_df.columns]
                            base_cols = ['date']
                            cols_to_export = base_cols + [c for c in cols_to_export if c not in base_cols]

                            if 'adj_prc' in group_df.columns and 'adj_prc' not in cols_to_export:
                                cols_to_export.append('adj_prc')
                            if 'ret' in group_df.columns and 'ret' not in cols_to_export:
                                cols_to_export.append('ret')

                            export_df = group_df[cols_to_export].sort_values('date').copy()

                            if (len(selected_financial_features) > 0 or enable_rolling_analysis) and price_wide is not None and ticker in price_wide.columns:
                                temp_fin = pd.DataFrame({'date': price_wide.index})
                                temp_fin['date'] = pd.to_datetime(temp_fin['date'])

                                if not returns_df.empty and ticker in returns_df.columns:
                                    temp_fin['daily_return_wrds_ret'] = returns_df[ticker].values
                                if 'log_return_py' in selected_financial_features and not log_returns_df.empty and ticker in log_returns_df.columns:
                                    temp_fin['log_return_py'] = log_returns_df[ticker].values
                                if 'cumulative_return_py' in selected_financial_features and not cumulative_return_df.empty and ticker in cumulative_return_df.columns:
                                    temp_fin['cumulative_return_py'] = cumulative_return_df[ticker].values
                                if 'volatility_annualized_py' in selected_financial_features and not rolling_vol_20_df.empty and ticker in rolling_vol_20_df.columns:
                                    temp_fin['rolling_volatility_20d_annualized_py'] = rolling_vol_20_df[ticker].values

                                if enable_rolling_analysis:
                                    if rolling_show_return and not rolling_mean_return_df.empty and ticker in rolling_mean_return_df.columns:
                                        temp_fin[f'rolling_mean_return_{rolling_window}d_py'] = rolling_mean_return_df[ticker].values
                                    if rolling_show_vol and not rolling_volatility_df.empty and ticker in rolling_volatility_df.columns:
                                        temp_fin[f'rolling_volatility_{rolling_window}d_py'] = rolling_volatility_df[ticker].values

                                export_df['date'] = pd.to_datetime(export_df['date'])
                                export_df = pd.merge(export_df, temp_fin, on='date', how='left')

                            if not summary_df.empty and 'Ticker' in summary_df.columns and ticker in summary_df['Ticker'].values:
                                row_summary = summary_df[summary_df['Ticker'] == ticker].copy()
                                for col in row_summary.columns:
                                    if col != 'Ticker':
                                        export_df[col] = row_summary.iloc[0][col]

                            if enable_rolling_analysis and not rolling_summary_df.empty and ticker in rolling_summary_df['Ticker'].values:
                                row_roll = rolling_summary_df[rolling_summary_df['Ticker'] == ticker].copy()
                                for col in row_roll.columns:
                                    if col != 'Ticker':
                                        export_df[col] = row_roll.iloc[0][col]

                            sheet_name = str(ticker)[:31]
                            export_df.to_excel(writer, index=False, sheet_name=sheet_name, startrow=1)
                            ws = writer.sheets[sheet_name]
                            format_worksheet(ws)

                    if "Summary Comparison" in selected_export_sections and not df_pivot.empty:
                        df_pivot.to_excel(writer, index=False, sheet_name="Summary_Comparison", startrow=1)
                        ws = writer.sheets["Summary_Comparison"]
                        format_worksheet(ws)

                    if "Raw WRDS Data" in selected_export_sections and not df_raw.empty:
                        df_raw.to_excel(writer, index=False, sheet_name="Raw_WRDS_Data", startrow=1)
                        ws = writer.sheets["Raw_WRDS_Data"]
                        format_worksheet(ws)

                    if "Enhanced Long Format" in selected_export_sections and not enhanced_long_df.empty:
                        enhanced_long_df.to_excel(writer, index=False, sheet_name="Enhanced_Long_Format", startrow=1)
                        ws = writer.sheets["Enhanced_Long_Format"]
                        format_worksheet(ws)

                    if "Financial Time Series" in selected_export_sections and not financial_timeseries_export.empty:
                        financial_timeseries_export.to_excel(writer, index=False, sheet_name="Financial_TimeSeries", startrow=1)
                        ws = writer.sheets["Financial_TimeSeries"]
                        format_worksheet(ws)

                    if "Financial Summary" in selected_export_sections and not summary_df.empty:
                        summary_df.to_excel(writer, index=False, sheet_name="Financial_Summary", startrow=1)
                        ws = writer.sheets["Financial_Summary"]
                        format_worksheet(ws)

                    if "Rolling Time Series" in selected_export_sections and enable_rolling_analysis and not rolling_export_df.empty:
                        rolling_export_df.to_excel(writer, index=False, sheet_name="Rolling_TimeSeries", startrow=1)
                        ws = writer.sheets["Rolling_TimeSeries"]
                        format_worksheet(ws)

                    if "Rolling Summary" in selected_export_sections and enable_rolling_analysis and not rolling_summary_df.empty:
                        rolling_summary_df.to_excel(writer, index=False, sheet_name="Rolling_Summary", startrow=1)
                        ws = writer.sheets["Rolling_Summary"]
                        format_worksheet(ws)

                    if "Descriptive Stats" in selected_export_sections and not descriptive_stats_df.empty:
                        descriptive_stats_df.to_excel(writer, index=False, sheet_name="Descriptive_Stats", startrow=1)
                        ws = writer.sheets["Descriptive_Stats"]
                        format_worksheet(ws)

                    if "Data Issues" in selected_export_sections and not data_issues_df.empty:
                        data_issues_df.to_excel(writer, index=False, sheet_name="Data_Issues", startrow=1)
                        ws = writer.sheets["Data_Issues"]
                        format_worksheet(ws)

                    if "Adjusted Prices" in selected_export_sections and not adjusted_price_export_df.empty:
                        adjusted_price_export_df.to_excel(writer, index=False, sheet_name="Adjusted_Prices", startrow=1)
                        ws = writer.sheets["Adjusted_Prices"]
                        format_worksheet(ws)

                    if "Correlation Matrix" in selected_export_sections and not correlation_matrix.empty:
                        correlation_matrix.reset_index().to_excel(writer, index=False, sheet_name="Correlation_Matrix", startrow=1)
                        ws = writer.sheets["Correlation_Matrix"]
                        format_worksheet(ws)

                    if "Metric Guide" in selected_export_sections:
                        guide_metric_df[show_cols].to_excel(writer, index=False, sheet_name="Metric_Guide", startrow=1)
                        ws = writer.sheets["Metric_Guide"]
                        format_worksheet(ws)

                    if "Charts" in selected_export_sections and include_chart_in_excel:
                        if raw_fig is not None:
                            ws_name = "Charts"
                            ws_chart = workbook[ws_name] if ws_name in workbook.sheetnames else workbook.create_sheet(title=ws_name)
                            safe_add_image_to_sheet(ws_chart, raw_fig, "A1", temp_files, width=800, height=400)

                        if compare_fig is not None:
                            ws_name = "Enhanced_Charts"
                            ws_compare = workbook[ws_name] if ws_name in workbook.sheetnames else workbook.create_sheet(title=ws_name)
                            safe_add_image_to_sheet(ws_compare, compare_fig, "A1", temp_files, width=800, height=420)

                        if cum_fig is not None:
                            ws_name = "Financial_Charts"
                            ws_fin = workbook[ws_name] if ws_name in workbook.sheetnames else workbook.create_sheet(title=ws_name)
                            safe_add_image_to_sheet(ws_fin, cum_fig, "A1", temp_files, width=800, height=400)

                        if heatmap_fig is not None:
                            ws_name = "Financial_Charts"
                            ws_fin = workbook[ws_name] if ws_name in workbook.sheetnames else workbook.create_sheet(title=ws_name)
                            safe_add_image_to_sheet(ws_fin, heatmap_fig, "A25", temp_files, width=700, height=500)

                        if enable_rolling_analysis and rolling_return_fig is not None:
                            ws_name = "Rolling_Charts"
                            ws_roll = workbook[ws_name] if ws_name in workbook.sheetnames else workbook.create_sheet(title=ws_name)
                            safe_add_image_to_sheet(ws_roll, rolling_return_fig, "A1", temp_files, width=800, height=400)

                        if enable_rolling_analysis and rolling_vol_fig is not None:
                            ws_name = "Rolling_Charts"
                            ws_roll = workbook[ws_name] if ws_name in workbook.sheetnames else workbook.create_sheet(title=ws_name)
                            safe_add_image_to_sheet(ws_roll, rolling_vol_fig, "A25", temp_files, width=800, height=400)

            finally:
                for tf in temp_files:
                    if os.path.exists(tf):
                        try:
                            os.remove(tf)
                        except:
                            pass

                output.seek(0)
                safe_export_name = sanitize_excel_filename(export_file_prefix)
                export_filename = f"{safe_export_name}.xlsx" if not safe_export_name.lower().endswith(".xlsx") else safe_export_name

                st.download_button(
                    label="Download Full Excel Report (.xlsx)",
                    data=output,
                    file_name=export_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"Processing Failed: {e}")
            import traceback
            st.error(traceback.format_exc())

# =========================================================
# MODULE 2: Report Generator (Enhanced with SEC EDGAR & DCF)
# =========================================================
elif APP_PAGE == "Report Generator":
    st.title("📄 SEC EDGAR Standardized Financial Report")
    st.caption("Generate investment reports in Word format with dynamic valuation and DuPont analytics.")

    if st.session_state.db_engine is None:
        st.warning("⚠️ Please connect to the WRDS Terminal in the sidebar first to fetch market data.")
    else:
        engine = st.session_state.db_engine

        # ==========================================
        # 新增第三个 TAB: 📈 Market Comparison
        # ==========================================
        tab_report, tab_dcf, tab_compare = st.tabs(["📝 Report Generator", "🧮 Interactive DCF Model", "📈 Market Comparison"])

        # -----------------------------------------------------
        # SUB-TAB 1: Report Generator
        # -----------------------------------------------------
        with tab_report:
            with st.form("rg_form"):
                st.header("1. Target Company Selection")
                col1, col2 = st.columns(2)
                with col1:
                    tickers_input_rg = st.text_input("Target Tickers (comma separated, e.g., AAPL, MSFT)", value="AAPL", placeholder="AAPL, MSFT").upper()
                with col2:
                    benchmark = st.selectbox("Market Benchmark (for relative performance)", ["SPY", "QQQ", "DIA"])

                st.header("2. Analysis Parameters")
                col3, col4, col5 = st.columns(3)
                with col3: start_date_rg = st.date_input("Start Date", value=datetime.date(2022, 1, 1))
                with col4: end_date_rg = st.date_input("End Date", value=datetime.date.today())
                with col5: rolling_window_rg = st.number_input("Rolling Window (Days)", min_value=20, value=60, step=10)

                gen_btn = st.form_submit_button("📝 Generate Investment Report", use_container_width=True)

            if gen_btn:
                if not tickers_input_rg.strip():
                    st.warning("⚠️ Target tickers cannot be empty! Please enter at least one ticker above (e.g., AAPL).")
                else:
                    try:
                        ticker_list_rg = [t.strip() for t in tickers_input_rg.split(",") if t.strip()]
                        all_tickers = ticker_list_rg + [benchmark]

                        with st.spinner("Retrieving Market Data (WRDS) & Financial Statements (SEC EDGAR)..."):
                            # --- 1. WRDS Market Data Fetching ---
                            t_str = "', '".join(all_tickers)
                            lookup_sql = f"SELECT DISTINCT permno, ticker FROM crsp.msenames WHERE ticker IN ('{t_str}') AND namedt <= '{end_date_rg}' AND nameendt >= '{start_date_rg}'"
                            
                            with engine.connect() as conn:
                                lookup_df = pd.read_sql(lookup_sql, conn)
                            
                            if not lookup_df.empty:
                                permno_str = ", ".join([str(int(x)) for x in lookup_df['permno'].unique()])
                                data_sql = f"SELECT date, permno, prc, ret, cfacpr FROM crsp.dsf WHERE permno IN ({permno_str}) AND date >= '{start_date_rg}' AND date <= '{end_date_rg}'"
                                with engine.connect() as conn:
                                    df_raw_rg = pd.read_sql(data_sql, conn)
                                
                                df_raw_rg['date'] = pd.to_datetime(df_raw_rg['date'])
                                df_raw_rg = pd.merge(df_raw_rg, lookup_df, on='permno', how='left').dropna(subset=['ticker'])
                                df_raw_rg['ticker'] = df_raw_rg['ticker'].str.upper()
                                df_raw_rg = apply_price_adjustment(df_raw_rg)

                                price_wide_rg = df_raw_rg.pivot_table(index='date', columns='ticker', values='adj_prc', aggfunc='last').ffill()
                                ret_wide_rg = df_raw_rg.pivot_table(index='date', columns='ticker', values='ret', aggfunc='last').fillna(0)
                                
                                bench_series = ret_wide_rg[benchmark] if benchmark in ret_wide_rg.columns else None
                                portfolio_prices = price_wide_rg.drop(columns=[benchmark], errors='ignore')
                                portfolio_rets = ret_wide_rg.drop(columns=[benchmark], errors='ignore')

                                rets, log_rets, cum_rets, summary_df_rg, corr_df, roll_vol = compute_financial_metrics(portfolio_prices, portfolio_rets, bench_series)
                            else:
                                st.warning("No market data found in WRDS. Report will only contain SEC EDGAR financials.")
                                cum_rets, roll_vol, corr_df, summary_df_rg = pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
                                price_wide_rg = pd.DataFrame()

                            # --- 2. SEC EDGAR API Integration ---
                            cik_map = get_sec_cik_mapping()
                            sec_financials = {}
                            sec_links = {}
                            sec_insights = {}
                            
                            for t in ticker_list_rg:
                                facts = fetch_sec_company_facts(t, cik_map)
                                if facts:
                                    sec_financials[t] = extract_financial_data(facts)
                                
                                links = fetch_sec_filings(t, cik_map, limit=8)
                                sec_links[t] = links
                                
                                if links:
                                    sec_insights[t] = fetch_sec_text_insights(links)

                            # --- 3. External Data Processing (Macro) ---
                            macro_df = fetch_world_bank_macro()

                        st.success("Data aggregation complete. Computing Advanced Financial Ratios & Valuations...")

                        # =========================================================
                        # PHASE 2: Data Depth & Valuation Calculation
                        # =========================================================
                        ratios_dict = {}
                        val_dict = {}
                        dupont_data = []

                        for t in ticker_list_rg:
                            f_data = sec_financials.get(t, {})
                            years = set()
                            for metric, dict_data in f_data.items():
                                years.update(dict_data.keys())
                            years = sorted(list(years))
                            
                            # Fetch WRDS Prices for Valuation & Technicals
                            latest_price = np.nan
                            high_52w = np.nan
                            low_52w = np.nan
                            sma50 = np.nan
                            sma200 = np.nan
                            
                            if not price_wide_rg.empty and t in price_wide_rg.columns:
                                series = price_wide_rg[t].dropna()
                                if not series.empty:
                                    latest_price = series.iloc[-1]
                                    last_252 = series.tail(252)
                                    high_52w = last_252.max()
                                    low_52w = last_252.min()
                                    if len(series) >= 50:
                                        sma50 = series.tail(50).mean()
                                    if len(series) >= 200:
                                        sma200 = series.tail(200).mean()

                            ticker_ratios = {}
                            latest_year = years[-1] if years else None
                            
                            for y in years:
                                rev = f_data.get('Total Revenue', {}).get(y)
                                gp = f_data.get('Gross Profit', {}).get(y)
                                ni = f_data.get('Net Income', {}).get(y)
                                ta = f_data.get('Total Assets', {}).get(y)
                                te = f_data.get('Stockholders Equity', {}).get(y)
                                eps = f_data.get('Earnings Per Share (EPS)', {}).get(y)
                                rnd = f_data.get('R&D Expense', {}).get(y)
                                sga = f_data.get('SG&A Expense', {}).get(y)
                                debt = f_data.get('Long-Term Debt', {}).get(y)
                                dps = f_data.get('Dividends Per Share', {}).get(y)
                                
                                gross_margin = (gp / rev) if (gp and rev and rev != 0) else np.nan
                                net_margin = (ni / rev) if (ni and rev and rev != 0) else np.nan
                                roe = (ni / te) if (ni and te and te != 0) else np.nan
                                roa = (ni / ta) if (ni and ta and ta != 0) else np.nan
                                rnd_margin = (rnd / rev) if (rnd and rev and rev != 0) else np.nan
                                sga_margin = (sga / rev) if (sga and rev and rev != 0) else np.nan
                                de_ratio = (debt / te) if (debt and te and te != 0) else np.nan
                                
                                ticker_ratios[y] = {
                                    'Gross Margin': gross_margin,
                                    'Net Margin': net_margin,
                                    'ROE': roe,
                                    'ROA': roa,
                                    'R&D Margin': rnd_margin,
                                    'SG&A Margin': sga_margin,
                                    'Debt/Equity Ratio': de_ratio
                                }

                                # Valuation & DuPont for the latest year
                                if y == latest_year:
                                    pe = (latest_price / eps) if pd.notna(latest_price) and eps and eps != 0 else np.nan
                                    shares_out = (ni / eps) if ni and eps and eps != 0 else np.nan
                                    pb = (latest_price / (te / shares_out)) if pd.notna(latest_price) and te and shares_out and shares_out != 0 else np.nan
                                    div_yield = (dps / latest_price) if pd.notna(latest_price) and pd.notna(dps) and latest_price != 0 else np.nan
                                    
                                    val_dict[t] = {
                                        'Latest Price': latest_price,
                                        '52W High': high_52w,
                                        '52W Low': low_52w,
                                        'P/E Ratio': pe,
                                        'P/B Ratio': pb,
                                        'Dividend Yield': div_yield,
                                        'SMA 50': sma50,
                                        'SMA 200': sma200
                                    }

                                    # DuPont Breakdown: Net Margin, Asset Turnover, Equity Multiplier
                                    asset_turnover = (rev / ta) if rev and ta and ta != 0 else 0
                                    equity_mult = (ta / te) if ta and te and te != 0 else 0
                                    dupont_data.append({
                                        'Ticker': t,
                                        'Net Margin': net_margin if pd.notna(net_margin) else 0,
                                        'Asset Turnover': asset_turnover,
                                        'Equity Multiplier': equity_mult
                                    })

                            # YoY Growth for latest year
                            yoy_rev, yoy_ni = np.nan, np.nan
                            if len(years) >= 2:
                                y1, y2 = years[-2], years[-1]
                                rev1, rev2 = f_data.get('Total Revenue', {}).get(y1), f_data.get('Total Revenue', {}).get(y2)
                                ni1, ni2 = f_data.get('Net Income', {}).get(y1), f_data.get('Net Income', {}).get(y2)
                                
                                yoy_rev = ((rev2 - rev1) / abs(rev1)) if rev1 and rev2 and rev1 != 0 else np.nan
                                yoy_ni = ((ni2 - ni1) / abs(ni1)) if ni1 and ni2 and ni1 != 0 else np.nan
                                ticker_ratios[latest_year]['YoY Revenue Growth'] = yoy_rev
                                ticker_ratios[latest_year]['YoY Net Income Growth'] = yoy_ni
                            
                            ratios_dict[t] = ticker_ratios

                        # --- Chart Generation (Fixed Dimensions) ---
                        px.defaults.template = "plotly_white"
                        charts = {}
                        if not cum_rets.empty:
                            fig_cum = px.line(cum_rets, title="Market Cumulative Returns (vs Benchmark)")
                            fig_cum.update_layout(xaxis_title="Date", yaxis_title="Return", width=900, height=450, margin=dict(l=20, r=20, t=40, b=20))
                            charts['Market Cumulative Returns'] = safe_plotly_to_png_bytes(fig_cum)

                        if not roll_vol.empty:
                            fig_vol = px.line(roll_vol, title=f"Annualized Rolling Volatility ({rolling_window_rg}-Day)")
                            fig_vol.update_layout(xaxis_title="Date", yaxis_title="Volatility", width=900, height=450, margin=dict(l=20, r=20, t=40, b=20))
                            charts['Historical Volatility'] = safe_plotly_to_png_bytes(fig_vol)

                        if not corr_df.empty:
                            fig_corr = px.imshow(corr_df, text_auto=".2f", color_continuous_scale='RdBu_r', title='Return Correlation Heatmap')
                            fig_corr.update_layout(width=900, height=500, margin=dict(l=20, r=20, t=40, b=20))
                            charts['Correlation Matrix'] = safe_plotly_to_png_bytes(fig_corr)

                        if dupont_data:
                            dupont_df = pd.DataFrame(dupont_data)
                            dupont_melt = dupont_df.melt(id_vars='Ticker', var_name='DuPont Component', value_name='Ratio')
                            fig_dupont = px.bar(dupont_melt, x='Ticker', y='Ratio', color='DuPont Component', barmode='group',
                                                title="DuPont Analysis Components (Latest FY)")
                            fig_dupont.update_layout(width=900, height=450, margin=dict(l=20, r=20, t=40, b=20))
                            charts['DuPont Analysis'] = safe_plotly_to_png_bytes(fig_dupont)

                        # =========================================================
                        # PHASE 3: Intelligent Narrative Generation
                        # =========================================================
                        narrative = []
                        
                        for t in ticker_list_rg:
                            v = val_dict.get(t, {})
                            r = ratios_dict.get(t, {}).get(latest_year, {})
                            
                            # Fetch VaR and MDD and Sharpe from summary
                            var95, mdd, sharpe = np.nan, np.nan, np.nan
                            if not summary_df_rg.empty and t in summary_df_rg['Ticker'].values:
                                row = summary_df_rg[summary_df_rg['Ticker'] == t].iloc[0]
                                var95 = row.get('VaR 95% (Historical)', np.nan)
                                mdd = row.get('Maximum Drawdown', np.nan)
                                sharpe = row.get('Sharpe Ratio', np.nan)

                            # Format strings
                            pe_str = f"{v.get('P/E Ratio', np.nan):.1f}x" if pd.notna(v.get('P/E Ratio')) else "N/A"
                            pb_str = f"{v.get('P/B Ratio', np.nan):.1f}x" if pd.notna(v.get('P/B Ratio')) else "N/A"
                            rev_g_str = format_percentage(r.get('YoY Revenue Growth', np.nan))
                            roe_str = format_percentage(r.get('ROE', np.nan))
                            var_str = format_percentage(var95)
                            mdd_str = format_percentage(mdd)
                            rnd_str = format_percentage(r.get('R&D Margin', np.nan))
                            de_str = format_float(r.get('Debt/Equity Ratio', np.nan)) + "x" if pd.notna(r.get('Debt/Equity Ratio', np.nan)) else "N/A"

                            # Technical & Sharpe Narrative
                            sma_text = ""
                            latest_p = v.get('Latest Price', np.nan)
                            sma200 = v.get('SMA 200', np.nan)
                            if pd.notna(latest_p) and pd.notna(sma200):
                                if latest_p > sma200:
                                    sma_text = f"Technically, the stock is trading at {format_currency(latest_p)}, above its 200-day moving average ({format_currency(sma200)}), indicating a bullish long-term trend. "
                                else:
                                    sma_text = f"Technically, the stock is trading at {format_currency(latest_p)}, below its 200-day moving average ({format_currency(sma200)}), suggesting a bearish long-term trend. "
                                    
                            sharpe_text = ""
                            if pd.notna(sharpe):
                                if sharpe > 1.0:
                                    sharpe_text = f"Its Sharpe Ratio of {sharpe:.2f} reflects strong risk-adjusted returns."
                                elif sharpe > 0:
                                    sharpe_text = f"Its Sharpe Ratio of {sharpe:.2f} reflects positive but moderate risk-adjusted returns."
                                else:
                                    sharpe_text = f"Its Sharpe Ratio of {sharpe:.2f} indicates negative risk-adjusted returns over the period."

                            # Construct Company Paragraph
                            n_text = f"[{t}] is currently trading at a P/E of {pe_str} and P/B of {pb_str}. {sma_text}"
                            n_text += f"Fundamentally, its latest YoY Revenue Growth stands at {rev_g_str} with an ROE of {roe_str}. "
                            if rnd_str != "N/A" or de_str != "N/A":
                                n_text += f"In terms of capital structure, the company maintains a Debt-to-Equity ratio of {de_str} and reinvests {rnd_str} of its revenue into R&D. "
                            n_text += f"On the risk front, its historical 95% VaR is {var_str}, and its maximum drawdown over the period was {mdd_str}. {sharpe_text}"
                            narrative.append(n_text)
                        
                        if not summary_df_rg.empty:
                            valid_summary = summary_df_rg.dropna(subset=['Annualized Return'])
                            if not valid_summary.empty:
                                best = valid_summary.loc[valid_summary['Annualized Return'].idxmax()]
                                narrative.append(f"Overall Conclusion: Relative to the {benchmark} benchmark, {best['Ticker']} exhibited the strongest annualized return profile ({best['Annualized Return']*100:.2f}%).")


                        # =========================================================
                        # PHASE 4 & 1: Document Construction & Formatting Excellence
                        # =========================================================
                        doc = Document()
                        
                        # --- Headers and Footers ---
                        section = doc.sections[0]
                        header = section.header
                        header_para = header.paragraphs[0]
                        header_para.text = f"Investment Research Report | {', '.join(ticker_list_rg)}"
                        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        footer = section.footer
                        footer_para = footer.paragraphs[0]
                        footer_para.text = f"Strictly Confidential | Generated by WRDS & SEC EDGAR"
                        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # --- Cover Page ---
                        doc.add_heading('INVESTMENT REPORT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                        
                        p_date = doc.add_paragraph()
                        p_date.add_run(f"Report Date: {datetime.date.today()}")
                        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        p_target = doc.add_paragraph()
                        p_target.add_run(f"Target Entities: {', '.join(ticker_list_rg)}")
                        p_target.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        p_bench = doc.add_paragraph()
                        p_bench.add_run(f"Market Benchmark: {benchmark}")
                        p_bench.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        doc.add_paragraph()
                        
                        # Executive Summary box on cover
                        doc.add_heading('Executive Summary & Investment Thesis', level=1)
                        if narrative:
                            for p_text in narrative:
                                doc.add_paragraph(p_text)
                        else:
                            doc.add_paragraph("Insufficient data to generate automated summary narrative.")

                        # --- Inject Business Overview and Quotes in Exec Summary ---
                        for t in ticker_list_rg:
                            if t in sec_insights:
                                doc.add_heading(f"{t} - Business Overview", level=3)
                                doc.add_paragraph(sec_insights[t]['Business Overview'])
                                doc.add_heading(f"{t} - Recent Management Quotes", level=3)
                                p_quote = doc.add_paragraph()
                                p_quote.add_run(sec_insights[t]['CEO Quotes']).italic = True

                        # --- Simplified Static TOC ---
                        add_chapter_heading(doc, 'Table of Contents', level=1)
                        toc_items = [
                            "Part I. Fundamental Analysis & Ratios",
                            "Part II. Market & Risk Analytics",
                            "Part III. Macroeconomic Data",
                            "Part IV. Appendix & Disclosures"
                        ]
                        for item in toc_items:
                            p = doc.add_paragraph(item)
                            p.paragraph_format.left_indent = Inches(0.5)

                        # --- PART I. Fundamental Analysis (SEC EDGAR) ---
                        add_chapter_heading(doc, 'Part I. Fundamental Analysis & Ratios', level=1)
                        doc.add_paragraph("Financial data is extracted via the SEC EDGAR XBRL API (US GAAP Form 10-K). Figures are in standard USD.")

                        for t in ticker_list_rg:
                            doc.add_heading(f"{t} - Financial Highlights", level=2)
                            fin_data = sec_financials.get(t, {})
                            
                            years = set()
                            for metric, dict_data in fin_data.items():
                                years.update(dict_data.keys())
                            years = sorted(list(years), reverse=True)[:3] 
                            
                            if years:
                                # 1. Absolute Financials Table
                                doc.add_heading("Core Financial Statements", level=3)
                                table = doc.add_table(rows=1, cols=len(years) + 1)
                                
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = "Financial Metric"
                                for i, y in enumerate(years):
                                    hdr_cells[i+1].text = f"FY {y}"
                                
                                display_metrics = [
                                    'Total Revenue', 'Gross Profit', 'Operating Income', 'Net Income',
                                    'Operating Cash Flow', 'Total Assets', 'Total Liabilities', 
                                    'Stockholders Equity', 'Long-Term Debt', 'R&D Expense', 'SG&A Expense',
                                    'Cash Dividends', 'Earnings Per Share (EPS)'
                                ]
                                for m in display_metrics:
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = m
                                    for i, y in enumerate(years):
                                        val = fin_data.get(m, {}).get(y)
                                        if m == 'Earnings Per Share (EPS)':
                                            row_cells[i+1].text = f"${val:.2f}" if pd.notna(val) else "N/A"
                                        else:
                                            row_cells[i+1].text = format_currency(val)
                                
                                set_wall_street_table_style(table)
                                doc.add_paragraph()

                                # 2. Financial Ratios Table
                                doc.add_heading("Key Ratios & Growth", level=3)
                                r_table = doc.add_table(rows=1, cols=len(years) + 1)
                                r_hdr = r_table.rows[0].cells
                                r_hdr[0].text = "Ratio / Metric"
                                for i, y in enumerate(years):
                                    r_hdr[i+1].text = f"FY {y}"
                                
                                ratio_keys = [
                                    'Gross Margin', 'Net Margin', 'ROE', 'ROA', 
                                    'R&D Margin', 'SG&A Margin', 'Debt/Equity Ratio',
                                    'YoY Revenue Growth', 'YoY Net Income Growth'
                                ]
                                for rk in ratio_keys:
                                    rc = r_table.add_row().cells
                                    rc[0].text = rk
                                    for i, y in enumerate(years):
                                        val = ratios_dict.get(t, {}).get(y, {}).get(rk, np.nan)
                                        if rk == 'Debt/Equity Ratio':
                                            rc[i+1].text = f"{val:.2f}x" if pd.notna(val) else "N/A"
                                        else:
                                            rc[i+1].text = format_percentage(val)
                                
                                set_wall_street_table_style(r_table)
                                doc.add_paragraph()
                                
                            else:
                                doc.add_paragraph("No standardized US GAAP financial data found in SEC EDGAR for this ticker.")
                            
                            # --- Inject MD&A ---
                            if t in sec_insights:
                                doc.add_heading(f"{t} - Management's Discussion & Analysis (Snippet)", level=3)
                                doc.add_paragraph(sec_insights[t]['MD&A'])
                        
                        if 'DuPont Analysis' in charts:
                            doc.add_heading('DuPont Analysis Breakdown', level=2)
                            # Width 6.0 Inches to strictly control page layout without spilling over
                            doc.add_picture(io.BytesIO(charts['DuPont Analysis']), width=Inches(6.0))
                            p_cap_dupont = doc.add_paragraph()
                            p_cap_dupont.add_run("Figure: ROE = Net Margin x Asset Turnover x Equity Multiplier.").italic = True

                        # --- PART II. Market & Risk Analytics (WRDS) ---
                        add_chapter_heading(doc, 'Part II. Market & Risk Analytics', level=1)
                        
                        # Valuation Overview
                        doc.add_heading('Valuation & Technical Overview', level=2)
                        val_df = pd.DataFrame(val_dict).T.reset_index().rename(columns={'index': 'Ticker'})
                        if not val_df.empty:
                            v_table = doc.add_table(rows=1, cols=len(val_df.columns))
                            for i, col in enumerate(val_df.columns):
                                v_table.rows[0].cells[i].text = str(col)
                            for _, row in val_df.iterrows():
                                rc = v_table.add_row().cells
                                for i, col in enumerate(val_df.columns):
                                    val = row[col]
                                    if pd.isna(val): rc[i].text = "N/A"
                                    elif col in ['P/E Ratio', 'P/B Ratio']: rc[i].text = f"{val:.2f}x"
                                    elif col == 'Dividend Yield': rc[i].text = format_percentage(val)
                                    elif col in ['Latest Price', '52W High', '52W Low', 'SMA 50', 'SMA 200']: rc[i].text = f"${val:.2f}"
                                    else: rc[i].text = str(val)
                            set_wall_street_table_style(v_table)
                            doc.add_paragraph()
                            
                            # --- Inject Forward Looking Statements near Valuation ---
                            for t in ticker_list_rg:
                                if t in sec_insights:
                                    doc.add_heading(f"{t} - Forward-Looking Guidance", level=3)
                                    doc.add_paragraph(sec_insights[t]['Forward Looking'])

                        # WRDS KPI (Transposed)
                        if not summary_df_rg.empty:
                            doc.add_heading('Key Performance Indicators (Transposed)', level=2)
                            clean_summary = summary_df_rg.dropna(axis=1, how='all').set_index('Ticker').T
                            table = doc.add_table(rows=1, cols=len(clean_summary.columns) + 1)
                            
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = "Analysis Metric"
                            for idx, ticker_name in enumerate(clean_summary.columns):
                                hdr_cells[idx+1].text = str(ticker_name)
                            
                            for metric, row in clean_summary.iterrows():
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(metric)
                                for idx, ticker_name in enumerate(clean_summary.columns):
                                    val = row[ticker_name]
                                    if pd.isna(val): 
                                        row_cells[idx+1].text = "N/A"
                                    elif isinstance(val, float) and abs(val) < 2 and any(k in str(metric) for k in ['Return', 'Volatility', 'Drawdown', 'VaR', 'CVaR']):
                                        row_cells[idx+1].text = format_percentage(val)
                                    elif isinstance(val, float): 
                                        row_cells[idx+1].text = format_float(val)
                                    else: 
                                        row_cells[idx+1].text = str(val)
                            set_wall_street_table_style(table)
                            doc.add_paragraph()
                            
                            # --- Inject Risk Factors below Quantitative Risks ---
                            for t in ticker_list_rg:
                                if t in sec_insights:
                                    doc.add_heading(f"{t} - Key Risk Factors (Snippet)", level=3)
                                    doc.add_paragraph(sec_insights[t]['Risk Factors'])

                        # Insert WRDS Charts
                        for title in ['Market Cumulative Returns', 'Historical Volatility', 'Correlation Matrix']:
                            if title in charts and charts[title]:
                                doc.add_heading(title, level=2)
                                doc.add_picture(io.BytesIO(charts[title]), width=Inches(6.0))
                                
                                p_cap_chart = doc.add_paragraph()
                                p_cap_chart.add_run(f"Figure: {title} derived from WRDS CRSP daily trading data.").italic = True

                        # --- PART III. Macro Data ---
                        add_chapter_heading(doc, 'Part III. Macroeconomic Data', level=1)
                        if not macro_df.empty:
                            summary_macro = macro_df.groupby(['Indicator', 'Country'])['Value'].mean().unstack().round(2)
                            doc.add_paragraph("Recent 5-Year World Bank Averages:")
                            m_table = doc.add_table(rows=1, cols=len(summary_macro.columns)+1)
                            m_table.rows[0].cells[0].text = "Indicator"
                            for idx, col in enumerate(summary_macro.columns): m_table.rows[0].cells[idx+1].text = str(col)
                            for idx, row in summary_macro.iterrows():
                                r_cells = m_table.add_row().cells
                                r_cells[0].text = str(idx)
                                for jdx, col in enumerate(summary_macro.columns): r_cells[jdx+1].text = str(row[col])
                            set_wall_street_table_style(m_table)
                            doc.add_paragraph()

                        # --- PART IV. Appendix & Disclosures ---
                        add_chapter_heading(doc, 'Part IV. Appendix & Disclosures', level=1)
                        doc.add_heading('Comprehensive SEC Filings Index', level=2)
                        doc.add_paragraph("This section includes a broadened scope of recent regulatory filings including Annual/Quarterly reports (10-K/10-Q), Material Events (8-K), Proxy Materials (DEF 14A), Specialized Disclosures (SD), and Prospectuses (424B2, FWP) indicating debt/preferred stock offerings.")
                        for t, filings in sec_links.items():
                            if filings:
                                doc.add_heading(f"{t} Corporate Filings", level=3)
                                for f in filings:
                                    p = doc.add_paragraph()
                                    p.add_run(f"[{f['Date']}] {f['Form']} ({f.get('Description', 'Filing')})").bold = True
                                    p.add_run(f"\nAccession: {f['Accession']}\nLink: {f['Link']}")
                        
                        doc.add_heading('Disclaimers', level=2)
                        doc.add_paragraph("This report is generated dynamically using quantitative data from WRDS CRSP and fundamental data from the SEC EDGAR API. "
                                          "It is provided for informational and analytical purposes only and does not constitute financial advice, an endorsement, "
                                          "or a recommendation to buy, hold, or sell any securities.")

                        # Save Doc
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                        # File Download
                        st.download_button(
                            label="📥 Download Investment Report (.docx)",
                            data=buffer,
                            file_name=f"Investment_Report_{datetime.date.today()}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                        # UI Preview
                        st.subheader("Report Generation Complete")
                        st.success("Investment Report has been generated successfully with dynamic valuation, text mining, and DuPont analytics!")

                    except Exception as e:
                        st.error(f"❌ Generation Error: {e}")
                        import traceback
                        st.error(traceback.format_exc())

        # -----------------------------------------------------
        # SUB-TAB 2: Interactive DCF Model
        # -----------------------------------------------------
        with tab_dcf:
            st.header("🧮 Interactive DCF Valuation Model")
            st.markdown("Estimate the intrinsic value of companies based on their Operating Cash Flow, projected growth, and discount rate.")

            with st.form("dcf_form"):
                dcf_col1, dcf_col2 = st.columns([1, 2])
                with dcf_col1:
                    tickers_input_dcf = st.text_input("Enter Tickers for DCF (comma separated)", "AAPL, MSFT", key="dcf_tick").upper()
                    dcf_g = st.slider("Expected Growth Rate (Next 5 Years)", min_value=0.0, max_value=0.50, value=0.10, step=0.01, format="%.2f")
                    dcf_wacc = st.slider("Discount Rate (WACC)", min_value=0.01, max_value=0.20, value=0.08, step=0.01, format="%.2f")
                    dcf_tg = st.slider("Terminal Growth Rate", min_value=0.0, max_value=0.05, value=0.02, step=0.005, format="%.3f")

                    update_dcf_btn = st.form_submit_button("🔄 Update DCF Valuation", use_container_width=True)

            if update_dcf_btn or st.session_state.get('dcf_initialized', False):
                st.session_state['dcf_initialized'] = True
                
                if tickers_input_dcf and engine:
                    dcf_tickers = [t.strip() for t in tickers_input_dcf.split(",") if t.strip()]
                    cik_map = get_sec_cik_mapping()
                    
                    dcf_results = []
                    dcf_charts = {}
                    dcf_metrics_ui = {}
                    
                    with st.spinner("Calculating DCF for selected tickers..."):
                        for tk in dcf_tickers:
                            facts = fetch_sec_company_facts(tk, cik_map)
                            if facts:
                                fin_data = extract_financial_data(facts)
                                years = set()
                                for k, v in fin_data.items(): years.update(v.keys())
                                years = sorted(list(years))
                                if years:
                                    latest_year = years[-1]
                                    ocf = fin_data.get('Operating Cash Flow', {}).get(latest_year, 0)
                                    debt = fin_data.get('Long-Term Debt', {}).get(latest_year, 0) or 0
                                    ni = fin_data.get('Net Income', {}).get(latest_year, 0)
                                    eps = fin_data.get('Earnings Per Share (EPS)', {}).get(latest_year, 0)

                                    current_price = get_latest_price(tk, engine)

                                    if ocf and pd.notna(eps) and eps != 0 and current_price and pd.notna(current_price):
                                        shares_out = ni / eps
                                        
                                        # Calculate DCF
                                        cfs = [ocf * (1 + dcf_g)**i for i in range(1, 6)]
                                        tv = cfs[-1] * (1 + dcf_tg) / (dcf_wacc - dcf_tg)

                                        pv_cfs = [cfs[i] / (1 + dcf_wacc)**(i+1) for i in range(5)]
                                        pv_tv = tv / (1 + dcf_wacc)**5

                                        ev = sum(pv_cfs) + pv_tv
                                        eq_value = ev - debt
                                        intrinsic_value = eq_value / shares_out

                                        diff = (intrinsic_value - current_price) / current_price
                                        status = "Undervalued" if diff > 0 else "Overvalued"
                                        
                                        dcf_results.append({
                                            "Ticker": tk,
                                            "Current Price": current_price,
                                            "Intrinsic Value": intrinsic_value,
                                            "Upside/Downside": diff,
                                            "Status": status
                                        })
                                        
                                        dcf_metrics_ui[tk] = {
                                            "current_price": current_price,
                                            "intrinsic_value": intrinsic_value,
                                            "diff": diff,
                                            "status": status,
                                            "ocf": ocf,
                                            "shares_out": shares_out,
                                            "debt": debt
                                        }

                                        # Waterfall Chart
                                        fig = go.Figure(go.Waterfall(
                                            name="DCF", orientation="v",
                                            measure=["relative", "relative", "relative", "relative", "relative", "relative", "total"],
                                            x=["Year 1 PV", "Year 2 PV", "Year 3 PV", "Year 4 PV", "Year 5 PV", "Terminal Value PV", "Enterprise Value"],
                                            textposition="outside",
                                            text=[format_currency(v) for v in pv_cfs] + [format_currency(pv_tv), format_currency(ev)],
                                            y=pv_cfs + [pv_tv, ev],
                                            connector={"line":{"color":"rgb(63, 63, 63)"}},
                                        ))
                                        fig.update_layout(
                                            title=f"{tk} Discounted Cash Flow (DCF) Waterfall",
                                            showlegend=False,
                                            template="plotly_dark",
                                            margin=dict(l=20, r=20, t=40, b=20)
                                        )
                                        dcf_charts[tk] = fig
                                    else:
                                        st.warning(f"{tk}: Incomplete data to perform DCF (missing OCF, EPS, or Price).")
                                else:
                                    st.warning(f"{tk}: No SEC financial data found.")
                            else:
                                st.warning(f"{tk}: Could not fetch SEC EDGAR facts.")
                    
                    if dcf_results:
                        if len(dcf_tickers) > 1:
                            # Tabs for Comparison and Individual Tickers
                            tabs_list = ["📊 Comparison Summary"] + [res["Ticker"] for res in dcf_results]
                            dcf_tabs = st.tabs(tabs_list)
                            
                            with dcf_tabs[0]:
                                st.subheader("DCF Valuation Comparison")
                                df_comp = pd.DataFrame(dcf_results)
                                
                                # Format for display
                                df_comp_display = df_comp.copy()
                                df_comp_display["Current Price"] = df_comp_display["Current Price"].apply(lambda x: f"${x:.2f}")
                                df_comp_display["Intrinsic Value"] = df_comp_display["Intrinsic Value"].apply(lambda x: f"${x:.2f}")
                                df_comp_display["Upside/Downside"] = df_comp_display["Upside/Downside"].apply(lambda x: f"{x*100:.2f}%")
                                st.dataframe(df_comp_display, use_container_width=True, hide_index=True)
                                
                                # Comparison Bar Chart
                                fig_comp = go.Figure()
                                fig_comp.add_trace(go.Bar(x=df_comp["Ticker"], y=df_comp["Current Price"], name='Current Price', marker_color='indianred'))
                                fig_comp.add_trace(go.Bar(x=df_comp["Ticker"], y=df_comp["Intrinsic Value"], name='Intrinsic Value', marker_color='lightseagreen'))
                                fig_comp.update_layout(
                                    title="Current Price vs. Intrinsic Value",
                                    xaxis_title="Ticker",
                                    yaxis_title="Price ($)",
                                    barmode='group',
                                    template="plotly_dark"
                                )
                                st.plotly_chart(fig_comp, use_container_width=True)
                                
                            # Populate individual tabs
                            for idx, res in enumerate(dcf_results):
                                tk = res["Ticker"]
                                with dcf_tabs[idx + 1]:
                                    metrics = dcf_metrics_ui[tk]
                                    color = "normal" if metrics["diff"] > 0 else "inverse"
                                    m1, m2, m3 = st.columns(3)
                                    m1.metric("Current Price", f"${metrics['current_price']:.2f}")
                                    m2.metric("Intrinsic Value", f"${metrics['intrinsic_value']:.2f}", f"{metrics['diff']*100:.2f}% ({metrics['status']})", delta_color=color)
                                    m3.metric("Base Operating Cash Flow", format_currency(metrics['ocf']))
                                    
                                    st.plotly_chart(dcf_charts[tk], use_container_width=True)
                                    st.caption(f"Assumptions for {tk}: Shares Outstanding = {metrics['shares_out']/1e9:.2f}B, Long-Term Debt = {format_currency(metrics['debt'])}")
                                    st.caption("*Note: Intrinsic Value is calculated as (Enterprise Value - Long Term Debt) / Shares Outstanding.")

                        else:
                            # Only one ticker, just display it normally without the "Comparison" tab
                            tk = dcf_results[0]["Ticker"]
                            metrics = dcf_metrics_ui[tk]
                            color = "normal" if metrics["diff"] > 0 else "inverse"
                            
                            st.subheader(f"{tk} Valuation Results")
                            m1, m2, m3 = st.columns(3)
                            m1.metric("Current Price", f"${metrics['current_price']:.2f}")
                            m2.metric("Intrinsic Value", f"${metrics['intrinsic_value']:.2f}", f"{metrics['diff']*100:.2f}% ({metrics['status']})", delta_color=color)
                            m3.metric("Base Operating Cash Flow", format_currency(metrics['ocf']))
                            
                            st.plotly_chart(dcf_charts[tk], use_container_width=True)
                            st.caption(f"Assumptions for {tk}: Shares Outstanding = {metrics['shares_out']/1e9:.2f}B, Long-Term Debt = {format_currency(metrics['debt'])}")
                            st.caption("*Note: Intrinsic Value is calculated as (Enterprise Value - Long Term Debt) / Shares Outstanding.")
                            
                elif not engine:
                    st.warning("Please connect to WRDS database in the sidebar.")
        
        # -----------------------------------------------------
        # SUB-TAB 3: Market Comparison (NEW FEATURE WITH EXPORT)
        # -----------------------------------------------------
        with tab_compare:
            st.header("📈 Market Comparison Analysis")
            st.markdown("Compare the performance and risk metrics of selected stocks against a market benchmark. Data is retrieved dynamically from WRDS.")

            with st.form("market_comp_form"):
                mc_col1, mc_col2 = st.columns(2)
                with mc_col1:
                    mc_tickers = st.text_input(
                        "Stock Tickers (comma separated)", 
                        value="AAPL, MSFT", 
                        help="Enter multiple stock symbols separated by commas"
                    ).upper()
                with mc_col2:
                    # UPDATED: Dropdown for benchmark selection
                    benchmark_options = {
                        "SPY": "SPY (S&P 500)",
                        "QQQ": "QQQ (Nasdaq 100)",
                        "DIA": "DIA (Dow Jones)",
                        "IWM": "IWM (Russell 2000)"
                    }
                    selected_bench_key = st.selectbox(
                        "Market Benchmark", 
                        options=list(benchmark_options.keys()), 
                        format_func=lambda x: benchmark_options[x],
                        help="Select a standard ETF to act as the market proxy for Alpha and Beta calculations."
                    )
                    mc_benchmark = selected_bench_key

                mc_col3, mc_col4 = st.columns(2)
                with mc_col3:
                    mc_start = st.date_input("Start Date", value=datetime.date(2023, 1, 1), key="mc_start")
                with mc_col4:
                    mc_end = st.date_input("End Date", value=datetime.date.today(), key="mc_end")

                mc_btn = st.form_submit_button("Compare Performance", use_container_width=True)

            if mc_btn or st.session_state.get('mc_initialized', False):
                if mc_btn: st.session_state['mc_initialized'] = True

                if not mc_tickers.strip() or not mc_benchmark.strip():
                    st.warning("⚠️ Please enter both stock tickers and a benchmark ticker.")
                elif st.session_state.db_engine is None:
                    st.warning("⚠️ Please connect to the WRDS Terminal in the sidebar first.")
                else:
                    stock_list = [t.strip() for t in mc_tickers.split(",") if t.strip()]
                    all_tks = stock_list + [mc_benchmark]

                    with st.spinner(f"Fetching Market Data from WRDS for {', '.join(all_tks)}..."):
                        t_str = "', '".join(all_tks)
                        lookup_sql = f"SELECT DISTINCT permno, ticker FROM crsp.msenames WHERE ticker IN ('{t_str}') AND namedt <= '{mc_end}' AND nameendt >= '{mc_start}'"

                        try:
                            with engine.connect() as conn:
                                lookup_df = pd.read_sql(lookup_sql, conn)

                            if not lookup_df.empty:
                                permno_str = ", ".join([str(int(x)) for x in lookup_df['permno'].unique()])
                                data_sql = f"SELECT date, permno, prc, ret, cfacpr FROM crsp.dsf WHERE permno IN ({permno_str}) AND date >= '{mc_start}' AND date <= '{mc_end}'"
                                
                                with engine.connect() as conn:
                                    df_raw_mc = pd.read_sql(data_sql, conn)

                                df_raw_mc['date'] = pd.to_datetime(df_raw_mc['date'])
                                df_raw_mc = pd.merge(df_raw_mc, lookup_df, on='permno', how='left').dropna(subset=['ticker'])
                                df_raw_mc['ticker'] = df_raw_mc['ticker'].str.upper()
                                df_raw_mc = apply_price_adjustment(df_raw_mc)

                                ret_wide_mc = df_raw_mc.pivot_table(index='date', columns='ticker', values='ret', aggfunc='last').fillna(0)

                                if mc_benchmark in ret_wide_mc.columns:
                                    bench_returns = ret_wide_mc[mc_benchmark]
                                    stock_returns = ret_wide_mc.drop(columns=[mc_benchmark], errors='ignore')

                                    # Calculate Cumulative Returns
                                    cum_rets_mc = (1 + ret_wide_mc).cumprod() - 1

                                    st.subheader("📈 Cumulative Return Comparison")
                                    fig_mc = px.line(cum_rets_mc, title=f"Stock Performance vs Benchmark ({mc_benchmark})")
                                    fig_mc.update_layout(xaxis_title="Date", yaxis_title="Cumulative Return", template="plotly_dark", hovermode="x unified")
                                    st.plotly_chart(fig_mc, use_container_width=True)

                                    st.subheader("🛡️ Risk & Return Profile")
                                    comp_rows = []
                                    trading_days = 252
                                    
                                    for tk in stock_list:
                                        if tk in stock_returns.columns:
                                            s_ret = stock_returns[tk]
                                            ann_ret = (1 + s_ret.mean())**trading_days - 1 if len(s_ret) > 0 else np.nan
                                            ann_vol = s_ret.std(ddof=1) * np.sqrt(trading_days) if len(s_ret) > 1 else np.nan
                                            sharpe = ann_ret / ann_vol if ann_vol else np.nan

                                            # Calculate Beta & Alpha & Additional Metrics
                                            merged = pd.concat([s_ret, bench_returns], axis=1).dropna()
                                            
                                            # MDD
                                            s_cum_local = (1 + s_ret).cumprod() - 1
                                            mdd = compute_max_drawdown(s_cum_local)
                                            
                                            if len(merged) > 1:
                                                cov = np.cov(merged.iloc[:, 0], merged.iloc[:, 1], ddof=1)[0, 1]
                                                var = np.var(merged.iloc[:, 1], ddof=1)
                                                beta = cov / var if var else np.nan
                                                
                                                ann_ret_b_local = (1 + bench_returns.mean())**trading_days - 1
                                                alpha = (ann_ret - (beta * ann_ret_b_local)) if pd.notna(beta) else np.nan
                                                
                                                # Correlation
                                                corr = merged.iloc[:, 0].corr(merged.iloc[:, 1])
                                                
                                                # TE & IR
                                                active_return = merged.iloc[:, 0] - merged.iloc[:, 1]
                                                tracking_error = active_return.std(ddof=1) * np.sqrt(trading_days)
                                                information_ratio = (ann_ret - ann_ret_b_local) / tracking_error if tracking_error and tracking_error != 0 else np.nan
                                            else:
                                                beta, alpha, corr, tracking_error, information_ratio = np.nan, np.nan, np.nan, np.nan, np.nan

                                            comp_rows.append({
                                                "Ticker": tk,
                                                "Annualized Return": format_percentage(ann_ret),
                                                "Annualized Volatility": format_percentage(ann_vol),
                                                "Sharpe Ratio": format_float(sharpe),
                                                "Max Drawdown": format_percentage(mdd),
                                                "Correlation vs Bench": format_float(corr),
                                                "Beta (vs Benchmark)": format_float(beta),
                                                "Alpha (Annualized)": format_percentage(alpha),
                                                "Tracking Error": format_percentage(tracking_error),
                                                "Information Ratio": format_float(information_ratio)
                                            })

                                    # Add benchmark row for reference
                                    ann_ret_b = (1 + bench_returns.mean())**trading_days - 1
                                    ann_vol_b = bench_returns.std(ddof=1) * np.sqrt(trading_days)
                                    sharpe_b = ann_ret_b / ann_vol_b if ann_vol_b else np.nan
                                    mdd_b = compute_max_drawdown((1 + bench_returns).cumprod() - 1)
                                    
                                    comp_rows.append({
                                        "Ticker": f"{mc_benchmark} (Benchmark)",
                                        "Annualized Return": format_percentage(ann_ret_b),
                                        "Annualized Volatility": format_percentage(ann_vol_b),
                                        "Sharpe Ratio": format_float(sharpe_b),
                                        "Max Drawdown": format_percentage(mdd_b),
                                        "Correlation vs Bench": "1.0000",
                                        "Beta (vs Benchmark)": "1.0000",
                                        "Alpha (Annualized)": "0.00%",
                                        "Tracking Error": "0.00%",
                                        "Information Ratio": "N/A"
                                    })

                                    # Display table
                                    df_metrics = pd.DataFrame(comp_rows)
                                    st.dataframe(df_metrics, use_container_width=True, hide_index=True)

                                    # ---  (Risk vs Return) ---
                                    st.subheader("⚖️ Risk vs Return Profile")
                                    # Copy and convert strings to float for plotting
                                    df_plot = df_metrics.copy()
                                    df_plot['Ret_Float'] = df_plot['Annualized Return'].str.rstrip('%').astype('float') / 100.0
                                    df_plot['Vol_Float'] = df_plot['Annualized Volatility'].str.rstrip('%').astype('float') / 100.0

                                    fig_scatter = px.scatter(
                                        df_plot, x="Vol_Float", y="Ret_Float", text="Ticker", color="Ticker",
                                        labels={"Vol_Float": "Annualized Volatility", "Ret_Float": "Annualized Return"},
                                        title="Risk vs. Return Assessment"
                                    )
                                    # Formating axes as percentage
                                    fig_scatter.layout.xaxis.tickformat = ',.1%'
                                    fig_scatter.layout.yaxis.tickformat = ',.1%'
                                    fig_scatter.update_traces(textposition='top center', marker=dict(size=12))
                                    fig_scatter.update_layout(template="plotly_dark", showlegend=False)
                                    st.plotly_chart(fig_scatter, use_container_width=True)

                                    # ---  (Return Distribution) ---
                                    st.subheader("📊 Return Distribution (Histogram)")
                                    fig_hist = go.Figure()
                                    # Add benchmark trace
                                    fig_hist.add_trace(go.Histogram(
                                        x=bench_returns, name=f"{mc_benchmark} (Bench)", 
                                        opacity=0.6, histnorm='probability', marker_color='gray'
                                    ))
                                    # Add stock traces
                                    colors = px.colors.qualitative.Plotly
                                    for i, tk in enumerate(stock_list):
                                        if tk in stock_returns.columns:
                                            fig_hist.add_trace(go.Histogram(
                                                x=stock_returns[tk], name=tk, 
                                                opacity=0.6, histnorm='probability', marker_color=colors[i % len(colors)]
                                            ))
                                    fig_hist.update_layout(
                                        barmode='overlay', 
                                        title="Daily Return Distribution Comparison", 
                                        xaxis_title="Daily Return", 
                                        yaxis_title="Probability", 
                                        template="plotly_dark"
                                    )
                                    st.plotly_chart(fig_hist, use_container_width=True)

                                    # ---  (Export) ---
                                    st.markdown("---")
                                    st.subheader("📥 Export Comparison Report")
                                    
                                    exp_col1, exp_col2 = st.columns([1, 2])
                                    with exp_col1:
                                        mc_export_name = st.text_input("Export File Name", value="Market_Comparison_Report")
                                    
                                    with exp_col2:
                                        st.markdown("<br>", unsafe_allow_html=True)
                                        output_mc = io.BytesIO()
                                        temp_files = []
                                        try:
                                            with pd.ExcelWriter(output_mc, engine='openpyxl') as writer:
                                                # Write Comparison Data
                                                df_metrics.to_excel(writer, index=False, sheet_name="Comparison_Metrics")
                                                format_worksheet(writer.sheets["Comparison_Metrics"])
                                                
                                                # Render and Add Charts
                                                ws_charts = writer.book.create_sheet(title="Charts")
                                                safe_add_image_to_sheet(ws_charts, fig_mc, "A1", temp_files)
                                                safe_add_image_to_sheet(ws_charts, fig_scatter, "A25", temp_files)
                                                safe_add_image_to_sheet(ws_charts, fig_hist, "A50", temp_files)
                                                
                                            output_mc.seek(0)
                                            
                                            st.download_button(
                                                label="📥 Download Excel Report (Data + Charts)",
                                                data=output_mc,
                                                file_name=f"{sanitize_excel_filename(mc_export_name)}.xlsx",
                                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                                use_container_width=True
                                            )
                                        except Exception as e:
                                            st.error(f"Failed to generate export file: {e}")
                                        finally:
                                            for tf in temp_files:
                                                if os.path.exists(tf):
                                                    try:
                                                        os.remove(tf)
                                                    except:
                                                        pass

                                else:
                                    st.error(f"Benchmark '{mc_benchmark}' not found in the retrieved WRDS data. Make sure it's a valid traded symbol.")
                            else:
                                st.error("No valid ticker mapping found for the requested symbols in the specified date range. Please verify the tickers.")
                        except Exception as e:
                            st.error(f"Error fetching data from WRDS: {e}")
